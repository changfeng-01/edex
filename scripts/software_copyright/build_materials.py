from __future__ import annotations

import ast
import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


FULL_NAME = "芯智调参：基于仿真数据的电路参数智能推荐系统"
SHORT_NAME = "芯智调参"
VERSION = "V1.0"
PACKAGE_VERSION = "1.0.0"
BASELINE_COMMIT = "4ea050e47957a3d984ca9ea95f71c53be0e0f3cf"
BRANCH = "codex/software-copyright-template-aligned"
COMPLETION_DATE = "2026年07月15日"
FIRST_PUBLICATION_DATE = "2026年07月15日"
FIRST_PUBLICATION = "GitHub 公开仓库"
BOUNDARIES = (
    "data_source = real_simulation_csv",
    "engineering_validity = simulation_only",
    "must_resimulate = true",
)
OFFICIAL_FORM_URL = (
    "https://banshi.beijing.gov.cn/pubtask/task/1/110000000000/"
    "3e283672-76be-4c8c-98e8-0bebe9bd06bf.html"
    "?locationCode=110000000000&serverType=1002"
)
REGULATION_URL = "https://www.ncac.gov.cn/xxfb/flfg/bmgz/202410/t20241015_869486.html"

ROOT = Path(__file__).resolve().parents[2]
DEFAULT_OUTPUT = Path(r"D:\EDA大赛\output\software_copyright\CircuitPilot_V1.0_template_aligned")
SCREENSHOT_SOURCE = ROOT / "output" / "playwright" / "software_copyright"
MANIFEST_SCHEMA_VERSION = "1.1"
TEMPLATE_ALIGNMENT = {
    "application_form": "3.申请表模板.doc",
    "source_code": "4.程序鉴别材料模板.doc",
    "manual_primary": "5.3说明书-虚拟机管理系统（应用软件）.doc",
    "manual_screenshots": "5.5说明书-网页截图软件（应用软件）.doc",
    "cooperation_contract": "6.技术开发（合作）合同模板.doc",
    "approved_example": "2023年8月审核通过的申请实例材料一套.zip",
}

DOCUMENT_SPECS = {
    "00_材料总目录与提交检查清单": {"pdf_pages": 3, "submission_role": "内部总控"},
    "01_软件著作权登记申请表填报底稿": {"pdf_pages": 6, "submission_role": "官方系统录入底稿"},
    "02_业务理解与软件设计说明书": {"pdf_pages": 12, "submission_role": "内部支撑材料"},
    "03_软件操作手册暨文档鉴别材料": {"pdf_pages": 18, "submission_role": "正式文档鉴别材料"},
    "04_源程序鉴别材料_前30页后30页": {"pdf_pages": 60, "submission_role": "正式源程序鉴别材料"},
    "05_技术开发（合作）合同": {"pdf_pages": 16, "submission_role": "权属证明材料"},
    "05A_共同开发与著作权归属确认书": {"pdf_pages": 4, "submission_role": "合同附件"},
    "06_第三方组件与权利边界说明": {"pdf_pages": 3, "submission_role": "内部支撑材料"},
}

BLACK = "000000"
GRAY = "666666"
LIGHT = "EDEDED"
WHITE = "FFFFFF"
ACCENT = "1F4E79"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def ensure_safe_output_path(output: Path) -> None:
    if output.name.casefold() == "circuitpilot_v1.0":
        raise RuntimeError("拒绝写入旧版材料目录 CircuitPilot_V1.0；请使用模板对齐版新目录。")
    if (
        output.is_dir()
        and next(output.iterdir(), None) is not None
        and os.environ.get("CIRCUITPILOT_ALLOW_MATERIAL_REBUILD") != "1"
    ):
        raise RuntimeError("模板对齐版输出目录非空；默认拒绝覆盖，请改用新目录或显式启用受控重建。")


def run_text(args: Sequence[str]) -> str:
    return subprocess.check_output(args, cwd=ROOT, text=True, encoding="utf-8").strip()


def read_baseline_blob(relative_path: str) -> bytes:
    return subprocess.check_output(
        ["git", "show", f"{BASELINE_COMMIT}:{relative_path}"],
        cwd=ROOT,
    )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: float = 10.5,
                 bold: bool | None = None, color: str = BLACK) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def configure_document(doc: Document, *, compact: bool = False) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.top_margin = Cm(1.8 if compact else 2.2)
    sec.bottom_margin = Cm(1.8 if compact else 2.2)
    sec.left_margin = Cm(1.8 if compact else 2.4)
    sec.right_margin = Cm(1.8 if compact else 2.4)
    sec.header_distance = Cm(0.7)
    sec.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.45
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size in [("Title", 24), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    set_header_footer(sec)
    doc.core_properties.title = FULL_NAME
    doc.core_properties.subject = f"计算机软件著作权登记申请材料 {VERSION}"
    doc.core_properties.author = "共同开发团队"
    doc.core_properties.last_modified_by = "共同开发团队"
    doc.core_properties.comments = "申报材料；不含个人身份信息。"
    doc.core_properties.keywords = "软件著作权, 电路仿真, 参数推荐, CircuitPilot"


def set_header_footer(section, header_text: str | None = None) -> None:
    header = section.header
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run(header_text or f"{FULL_NAME}  {VERSION}")
    set_run_font(run, east_asia="宋体", size=8.5, color=GRAY)
    fp = section.footer.paragraphs[0]
    fp.clear()
    add_page_number(fp)
    for run in fp.runs:
        set_run_font(run, size=8.5, color=GRAY)


def set_fixed_page_header(section, left_text: str, page_no: int, *, font_size: float) -> None:
    header = section.header
    for child in list(header._element):
        header._element.remove(child)
    table = header.add_table(rows=1, cols=2, width=Cm(18.8))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(17.7)
    table.columns[1].width = Cm(1.1)
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.width = Cm(17.7)
    right.width = Cm(1.1)
    for cell in (left, right):
        cell.margin_top = 0
        cell.margin_bottom = 0
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(font_size + 1.5)
    left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = left.paragraphs[0].add_run(left_text)
    set_run_font(r, east_asia="宋体", latin="Arial Narrow", size=font_size)
    r._element.get_or_add_rPr().append(_char_scale(75))
    r = right.paragraphs[0].add_run(str(page_no))
    set_run_font(r, east_asia="宋体", latin="Arial", size=font_size + 1, bold=True)
    trailing = header.add_paragraph()
    trailing.paragraph_format.space_before = Pt(0)
    trailing.paragraph_format.space_after = Pt(0)
    trailing.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    trailing.paragraph_format.line_spacing = Pt(1)


def set_source_page_header(section, path_text: str, page_no: int) -> None:
    header = section.header
    for child in list(header._element):
        header._element.remove(child)
    table = header.add_table(rows=1, cols=3, width=Cm(14.75))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = (8.2, 2.3, 4.25)
    values = (f"{FULL_NAME}{VERSION}", "源代码", f"第{page_no}页，共60页")
    alignments = (WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT)
    for index, (width, value, alignment) in enumerate(zip(widths, values, alignments)):
        cell = table.cell(0, index)
        cell.width = Cm(width)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(value)
        set_run_font(run, east_asia="宋体", latin="Times New Roman", size=6.8, bold=(index == 1))
        if index == 0:
            run._element.get_or_add_rPr().append(_char_scale(70))
    path_paragraph = header.add_paragraph()
    path_paragraph.paragraph_format.space_before = Pt(0)
    path_paragraph.paragraph_format.space_after = Pt(0)
    path_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    path_paragraph.paragraph_format.line_spacing = Pt(8)
    run = path_paragraph.add_run(f"文件路径：{path_text}")
    set_run_font(run, east_asia="宋体", latin="Courier New", size=6.2, color=GRAY)
    run._element.get_or_add_rPr().append(_char_scale(75))


def add_cover(doc: Document, title: str, subtitle: str, kind: str = "申报材料") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    r = p.add_run(FULL_NAME)
    set_run_font(r, east_asia="黑体", size=20, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run(title)
    set_run_font(r, east_asia="黑体", size=25, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(subtitle)
    set_run_font(r, east_asia="宋体", size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    for line in (f"软件简称：{SHORT_NAME}", f"版本号：{VERSION}", f"材料性质：{kind}", "申请方式：实际开发成员共同申请"):
        r = p.add_run(line + "\n")
        set_run_font(r, east_asia="宋体", size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(45)
    r = p.add_run("编制日期：2026年07月")
    set_run_font(r, size=10.5)
    doc.add_page_break()


def add_notice(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    p = cell.paragraphs[0]
    r = p.add_run(title + "：")
    set_run_font(r, east_asia="黑体", size=10.5, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    doc.add_paragraph()


def add_kv_table(doc: Document, rows: Sequence[tuple[str, str]], widths: tuple[float, float] = (4.0, 12.0)) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(widths[0])
        cells[1].width = Cm(widths[1])
        set_cell_shading(cells[0], LIGHT)
        for i, text in enumerate((key, value)):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, east_asia="黑体" if i == 0 else "宋体", size=9.5, bold=(i == 0))


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
              font_size: float = 8.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], LIGHT)
        p = table.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, east_asia="黑体", size=font_size, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.keep_together = True
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.keep_together = True
        p.add_run(item)


def add_boundary(doc: Document) -> None:
    add_notice(
        doc,
        "证据边界",
        "；".join(BOUNDARIES)
        + "。系统处理的是仿真 CSV 与可追溯工程工件，不代表物理样片、流片、实验室实测或性能保证。",
    )


def save_doc(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.created = datetime(2026, 7, 15, tzinfo=timezone.utc)
    doc.core_properties.modified = datetime(2026, 7, 16, tzinfo=timezone.utc)
    doc.save(path)


def source_group(path: str) -> tuple[int, str]:
    groups = [
        ("src/goa_eval/product_api/", 0, "产品 API"),
        ("src/goa_eval/product/", 1, "产品服务"),
        ("src/goa_eval/web_api/", 2, "只读仪表盘 API"),
        ("src/goa_eval/web/", 3, "上传分析服务"),
        ("frontend/src/", 4, "前端界面"),
        ("scripts/", 5, "工程脚本"),
        ("src/goa_eval/ca/", 6, "电路分析"),
        ("src/goa_eval/pia/", 7, "PIA 与候选生成"),
        ("src/goa_eval/llso/", 8, "LLSO 优化"),
        ("src/goa_eval/multi_agent/", 9, "多智能体协作"),
        ("src/goa_eval/", 10, "评估内核与命令行"),
    ]
    for prefix, rank, label in groups:
        if path.startswith(prefix):
            return rank, label
    return 99, "其他"


def language_for(path: str) -> str:
    return {".py": "Python", ".ts": "TypeScript", ".tsx": "TypeScript/TSX", ".css": "CSS"}[Path(path).suffix.lower()]


def list_source_files() -> list[dict]:
    tracked = run_text(["git", "ls-tree", "-r", "--name-only", BASELINE_COMMIT]).splitlines()
    accepted: list[dict] = []
    allowed_ext = {".py", ".ts", ".tsx", ".css"}
    for rel in tracked:
        posix = rel.replace("\\", "/")
        suffix = Path(posix).suffix.lower()
        if suffix not in allowed_ext:
            continue
        if not (posix.startswith("src/goa_eval/") or posix.startswith("frontend/src/") or posix.startswith("scripts/")):
            continue
        lower = posix.lower()
        parts = lower.split("/")
        if (
            "tests" in parts
            or "test" in parts
            or ".test." in lower
            or lower.endswith("_test.py")
            or "/__pycache__/" in lower
            or posix.startswith("scripts/software_copyright/")
        ):
            continue
        raw = read_baseline_blob(posix)
        text = raw.decode("utf-8", errors="replace")
        lines = text.splitlines()
        rank, group = source_group(posix)
        accepted.append(
            {
                "path": posix,
                "language": language_for(posix),
                "line_count": len(lines),
                "sha256": hashlib.sha256(raw).hexdigest(),
                "source_group": group,
                "_rank": rank,
                "_lines": lines,
            }
        )
    accepted.sort(key=lambda x: (x["_rank"], x["path"].lower()))
    return accepted


def write_source_manifest(output: Path, entries: list[dict]) -> None:
    with (output / "source_manifest.csv").open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=["order", "path", "language", "line_count", "sha256", "source_group"],
        )
        writer.writeheader()
        for i, item in enumerate(entries, 1):
            writer.writerow(
                {
                    "order": i,
                    "path": item["path"],
                    "language": item["language"],
                    "line_count": item["line_count"],
                    "sha256": item["sha256"],
                    "source_group": item["source_group"],
                }
            )


def flatten_source(entries: list[dict]) -> list[dict]:
    stream: list[dict] = []
    global_line = 0
    for order, item in enumerate(entries, 1):
        for file_line, text in enumerate(item["_lines"], 1):
            global_line += 1
            stream.append(
                {
                    "global_line": global_line,
                    "file_order": order,
                    "path": item["path"],
                    "file_line": file_line,
                    "text": text,
                }
            )
    return stream


def select_source_pages(stream: Sequence[dict]) -> list[list[dict]]:
    lines_per_page = 50
    segment_lines = 30 * lines_per_page
    if len(stream) < segment_lines * 2:
        raise RuntimeError("源程序不足以提取前后各30页。")
    selected = list(stream[:segment_lines]) + list(stream[-segment_lines:])
    numbered = [dict(item, display_line=index) for index, item in enumerate(selected, 1)]
    return [
        numbered[start : start + lines_per_page]
        for start in range(0, len(numbered), lines_per_page)
    ]


def add_source_page(doc: Document, page_lines: Sequence[dict], display_page: int) -> None:
    if display_page == 1:
        section = doc.sections[0]
    else:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.75)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.0)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    unique_paths: list[str] = []
    for line in page_lines:
        if line["path"] not in unique_paths:
            unique_paths.append(line["path"])
    path_text = unique_paths[0] if len(unique_paths) == 1 else f"{unique_paths[0]} → {unique_paths[-1]}"
    set_source_page_header(section, path_text, display_page)
    fp = section.footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(f"{display_page}/60")
    set_run_font(r, east_asia="宋体", size=8, color=BLACK)

    for item in page_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(10.2)
        p.paragraph_format.keep_together = True
        number_run = p.add_run(f"{item['display_line']:04d}  ")
        set_run_font(number_run, east_asia="宋体", latin="Courier New", size=7.2, color=GRAY)
        text = item["text"].expandtabs(4)
        code_run = p.add_run(text)
        set_run_font(code_run, east_asia="宋体", latin="Courier New", size=7.5)
        scale = source_char_scale(len(text))
        code_run._element.get_or_add_rPr().append(_char_scale(scale))


def source_char_scale(text_length: int) -> int:
    """Return a conservative Word character scale that prevents code wrapping."""
    return min(100, max(26, int(6500 / max(1, text_length))))


def _char_scale(value: int) -> OxmlElement:
    scale = OxmlElement("w:w")
    scale.set(qn("w:val"), str(value))
    return scale


def build_source_doc(output: Path, entries: list[dict]) -> dict:
    stream = flatten_source(entries)
    needed = 30 * 50
    page_groups = select_source_pages(stream)
    selected = [line for page in page_groups for line in page]
    doc = Document()
    doc._body.clear_content()
    configure_document(doc, compact=True)
    page_map = []
    for page_index, page_lines in enumerate(page_groups):
        add_source_page(doc, page_lines, page_index + 1)
        page_map.append(
            {
                "display_page": page_index + 1,
                "segment": "front" if page_index < 30 else "back",
                "display_start": page_lines[0]["display_line"],
                "display_end": page_lines[-1]["display_line"],
                "original_global_start": page_lines[0]["global_line"],
                "original_global_end": page_lines[-1]["global_line"],
                "global_start": page_lines[0]["global_line"],
                "global_end": page_lines[-1]["global_line"],
                "source_line_count": 50,
                "first_path": page_lines[0]["path"],
                "last_path": page_lines[-1]["path"],
            }
        )
    path = output / "04_源程序鉴别材料_前30页后30页.docx"
    save_doc(doc, path)
    support_dir = output / "internal_support"
    support_dir.mkdir(parents=True, exist_ok=True)
    digest_by_path = {item["path"]: item["sha256"] for item in entries}
    line_manifest = support_dir / "source_line_manifest.csv"
    with line_manifest.open("w", newline="", encoding="utf-8-sig") as file:
        writer = csv.DictWriter(
            file,
            fieldnames=[
                "submission_line",
                "display_page",
                "segment",
                "original_global_line",
                "path",
                "file_line",
                "file_sha256",
            ],
        )
        writer.writeheader()
        for item in selected:
            writer.writerow(
                {
                    "submission_line": item["display_line"],
                    "display_page": (item["display_line"] - 1) // 50 + 1,
                    "segment": "front" if item["display_line"] <= 1500 else "back",
                    "original_global_line": item["global_line"],
                    "path": item["path"],
                    "file_line": item["file_line"],
                    "file_sha256": digest_by_path[item["path"]],
                }
            )
    return {
        "path": path.name,
        "full_stream_lines": len(stream),
        "selected_lines": len(selected),
        "front_range": [1, needed],
        "back_range": [len(stream) - needed + 1, len(stream)],
        "pages": page_map,
        "line_manifest": "internal_support/source_line_manifest.csv",
    }


def inventory_design_lines(entries: list[dict]) -> list[str]:
    lines = [
        f"软件全称：{FULL_NAME}。",
        f"软件简称：{SHORT_NAME}；版本号：{VERSION}。",
        "软件定位：面向电路仿真数据的版本化分析、约束评估、候选参数管理和重仿真闭环。",
        "业务原则：候选参数是下一轮仿真输入，不是已经获得验证的工程结论。",
        *[f"证据边界：{item}。" for item in BOUNDARIES],
        "业务角色：工程操作员负责输入数据、外部仿真与结果回填。",
        "业务角色：评审人员读取约束、证据、候选、版本比较和交付报告。",
        "业务角色：共同开发成员维护源码、接口契约、测试和材料。",
        "系统架构：React 前端通过 Product API 访问项目化领域服务。",
        "系统架构：FastAPI 路由层将请求交给项目、输入、分析、实验、任务和比较服务。",
        "系统架构：SQLAlchemy 仓储保存结构化状态，本地工件库保存 CSV、JSON、YAML 与报告。",
        "核心流程：工作区→项目→设计版本→输入预览→分析运行→候选审批→仿真任务→结果回填→版本比较。",
        "状态约束：系统不会自动把候选状态转换成已验证结论。",
        "异常策略：输入缺失、格式不符、校验失败、状态冲突和工件缺失均返回结构化错误。",
        "安全边界：材料不包含密钥、环境变量、个人身份信息、私有数据或本地调试路径。",
    ]
    route_patterns = (
        re.compile(r"@router\.(get|post|put|patch|delete)\(\s*[rf]?['\"]([^'\"]+)"),
        re.compile(r"@app\.(get|post|put|patch|delete)\(\s*[rf]?['\"]([^'\"]+)"),
    )
    for order, item in enumerate(entries, 1):
        path = item["path"]
        lines.append(
            f"模块清单{order:03d}：{path}，分组为{item['source_group']}，"
            f"共{item['line_count']}行，摘要校验值{item['sha256'][:16]}。"
        )
        text = "\n".join(item["_lines"])
        if path.endswith(".py"):
            try:
                tree = ast.parse(text)
                for node in ast.walk(tree):
                    if isinstance(node, ast.ClassDef):
                        lines.append(f"类职责：{path} 定义 {node.name}，用于封装该模块的领域数据或行为。")
                    elif isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                        args = [arg.arg for arg in node.args.args[:6]]
                        kind = "异步接口" if isinstance(node, ast.AsyncFunctionDef) else "函数"
                        lines.append(f"{kind}清单：{path}::{node.name}({', '.join(args)})，实现可定位的模块职责。")
            except SyntaxError:
                lines.append(f"解析说明：{path} 按文本纳入鉴别材料，未进行语义改写。")
        else:
            for match in re.finditer(
                r"\b(?:export\s+)?(?:default\s+)?(?:async\s+)?(?:function|class|interface|type|const)\s+([A-Za-z_$][\w$]*)",
                text,
            ):
                lines.append(f"前端符号：{path} 定义 {match.group(1)}，用于页面、组件、类型或客户端逻辑。")
        for pattern in route_patterns:
            for method, route in pattern.findall(text):
                lines.append(f"接口契约：{method.upper()} {route}，路由实现位于 {path}。")
        imports = re.findall(r"^(?:from\s+([\w.]+)\s+import|import\s+([\w.]+))", text, flags=re.MULTILINE)
        for left, right in imports[:12]:
            dep = left or right
            lines.append(f"模块依赖：{path} 引用 {dep}，由依赖注入或显式调用形成可追溯关系。")
    while len(lines) < 2200:
        idx = len(lines) % len(entries)
        item = entries[idx]
        lines.append(
            f"实现核对：{item['path']} 属于{item['source_group']}，"
            f"其原始文件 SHA-256 为 {item['sha256']}，材料未改写源文件。"
        )
    return lines


def add_fixed_text_page(doc: Document, page_lines: Sequence[str], page_no: int, segment: str) -> None:
    if page_no == 1:
        section = doc.sections[0]
    else:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.45)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    set_fixed_page_header(
        section,
        f"{FULL_NAME}  {VERSION}  文档鉴别材料（{segment}）",
        page_no,
        font_size=6.5,
    )
    fp = section.footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("设计说明文档鉴别材料｜30行/页｜" + BASELINE_COMMIT[:12])
    set_run_font(r, size=7, color=GRAY)
    for line in page_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(16.0)
        p.paragraph_format.keep_together = True
        r = p.add_run(line)
        set_run_font(r, east_asia="宋体", size=7.4)


def build_document_identification(output: Path, entries: list[dict]) -> dict:
    stream = inventory_design_lines(entries)
    needed = 30 * 30
    selected = stream[:needed] + stream[-needed:]
    doc = Document()
    doc._body.clear_content()
    configure_document(doc, compact=True)
    pages = []
    for page_index in range(60):
        page_lines = selected[page_index * 30 : (page_index + 1) * 30]
        segment = "前30页" if page_index < 30 else "后30页"
        add_fixed_text_page(doc, page_lines, page_index + 1, segment)
        pages.append(
            {
                "display_page": page_index + 1,
                "segment": "front" if page_index < 30 else "back",
                "master_start": page_index * 30 + 1 if page_index < 30 else len(stream) - needed + (page_index - 30) * 30 + 1,
                "line_count": 30,
            }
        )
    path = output / "02A_文档鉴别材料_前30页后30页.docx"
    save_doc(doc, path)
    return {
        "path": path.name,
        "master_line_count": len(stream),
        "selected_line_count": len(selected),
        "front_range": [1, needed],
        "back_range": [len(stream) - needed + 1, len(stream)],
        "pages": pages,
    }


def build_catalog(output: Path, source_stats: dict) -> Path:
    doc = Document()
    configure_document(doc)
    add_cover(doc, "材料总目录与提交检查清单", "共同申请团队内部定稿底稿", "总目录与提交核对")
    doc.add_heading("一、材料包使用说明", level=1)
    add_notice(
        doc,
        "重要",
        "本材料包是登记准备底稿。正式《计算机软件著作权登记申请表》必须在中国版权保护中心登记系统中"
        "录入后在线生成并签章，不得用本包中的自制底稿替代或改动官方版式。",
    )
    add_kv_table(
        doc,
        [
            ("软件全称", FULL_NAME),
            ("软件简称", SHORT_NAME),
            ("版本号", VERSION),
            ("开发完成日期", COMPLETION_DATE),
            ("发表状态", "已发表"),
            ("首次发表日期/方式", f"{FIRST_PUBLICATION_DATE}；{FIRST_PUBLICATION}"),
            ("申请口径", "实际开发成员共同申请；合作开发；原始取得；全部权利"),
            ("申报基线", BASELINE_COMMIT),
            ("源码统计", f"{source_stats['files']}个自研源文件，{source_stats['lines']}行"),
        ],
    )
    doc.add_heading("二、正式提交与内部支撑目录", level=1)
    rows = [
        ("01", "官方系统在线生成的登记申请表", "正式提交", "全体申请人按系统要求签章；本包01仅作录入底稿"),
        ("02", "源程序鉴别材料前30页后30页", "正式提交", "提交04文件；核对60页、页码、版本和源码连续性"),
        ("03", "文档鉴别材料前30页后30页", "正式提交", "提交02A文件；正文每页30行"),
        ("04", "合作开发与著作权归属确认书", "建议正式提交/留档", "所有实际开发成员逐人签字"),
        ("05", "共同著作权人身份证明", "正式提交", "按登记系统要求准备清晰有效材料"),
        ("06", "业务理解与软件设计说明书", "内部支撑/备查", "用于解释架构、业务、API和边界"),
        ("07", "软件操作手册", "内部支撑/可作鉴别文档", "截图均来自 test_only 确定性演示环境"),
        ("08", "第三方组件与权利边界说明", "内部支撑/备查", "不主张第三方组件著作权"),
        ("09", "source_manifest.csv", "内部支撑", "记录源文件顺序、行数和 SHA-256"),
        ("10", "snapshot_manifest.json", "内部支撑", "记录基线、选择区间、文件摘要和验收结果"),
    ]
    add_table(doc, ["序号", "材料", "用途", "提交/核对要求"], rows, 8.2)
    doc.add_heading("三、允许保留的待填写字段", level=1)
    add_bullets(
        doc,
        [
            "共同著作权人姓名、证件类型、证件号码。",
            "共同著作权人联系地址、电话、电子邮箱。",
            "共同著作权人实际贡献描述与签名/签章。",
            "申请代表姓名、联系方式、授权签字。",
            "首次发表城市（如登记系统要求精确到城市）。",
        ],
    )
    doc.add_heading("四、共同著作权人资格核对", level=1)
    add_bullets(
        doc,
        [
            "仅列入对软件源代码、界面或技术文档作出可识别实际开发贡献的成员。",
            "非开发参赛成员、行政协调人员和指导教师不因参赛或指导身份自动成为共同著作权人。",
            "每位申请人的贡献描述应与代码提交、设计记录、任务分工或其他可核验证据一致。",
            "如存在单位任务、职务开发、外包、委托或既有协议，应在提交前由专业人员复核权属。",
        ],
    )
    doc.add_heading("五、提交前总检查", level=1)
    checks = [
        "已在官方系统录入本包01底稿并在线生成正式申请表。",
        "软件全称、简称、版本号在申请表、源码页、文档页和权属文件中完全一致。",
        "开发完成日期、首次发表日期和发表状态与可保存证据一致。",
        "所有实际开发成员均被逐行列明，身份材料完整，签字/签章无遗漏。",
        "源码鉴别材料恰为60页，每页50行真实源码；前后段选择范围与清单一致。",
        "文档鉴别材料恰为60页，每页30行正文；前后段选择范围与清单一致。",
        "未包含密码、密钥、环境变量、身份证件图像、私人数据或未说明的占位符。",
        "未宣称流片、芯片实测、大量实验验证或保证优化效果。",
        "第三方组件未被计入团队原创权利主张。",
        "已保存 GitHub 首次发表证据：仓库页面、提交/发布记录、日期和访问时间。",
    ]
    for item in checks:
        p = doc.add_paragraph()
        p.add_run("□ ").bold = True
        p.add_run(item)
    doc.add_heading("六、规范依据与风险提示", level=1)
    doc.add_paragraph(f"办理规范：{OFFICIAL_FORM_URL}")
    doc.add_paragraph(f"《计算机软件著作权登记办法》相关页面：{REGULATION_URL}")
    doc.add_paragraph(
        "本包根据仓库现状和用户给定申报口径生成，不构成法律意见。正式提交前应核对中国版权保护中心"
        "最新系统字段、签章要求、材料份数和地方办理差异。"
    )
    path = output / "00_材料总目录与提交检查清单.docx"
    save_doc(doc, path)
    return path


def build_application_draft(output: Path, source_stats: dict) -> Path:
    doc = Document()
    configure_document(doc)
    add_cover(doc, "软件著作权登记申请表填报底稿", "用于官方系统逐项录入，不替代正式申请表", "申请表录入底稿")
    add_notice(
        doc,
        "正式表生成要求",
        "必须将本底稿内容录入中国版权保护中心系统，由系统在线生成正式申请表并按要求签章。"
        "不得把本文件直接作为官方申请表提交，也不得修改官方表格版式。",
    )
    doc.add_heading("一、软件基本信息", level=1)
    add_kv_table(
        doc,
        [
            ("软件全称", FULL_NAME),
            ("软件简称", SHORT_NAME),
            ("版本号", VERSION),
            ("软件分类", "应用软件"),
            ("作品说明", "原创"),
            ("开发完成日期", COMPLETION_DATE),
            ("发表状态", "已发表"),
            ("首次发表日期", FIRST_PUBLICATION_DATE),
            ("首次发表地点", "中国【首次发表城市】"),
            ("首次发表方式", FIRST_PUBLICATION),
        ],
    )
    doc.add_heading("二、开发与权利信息", level=1)
    add_kv_table(
        doc,
        [
            ("开发方式", "合作开发"),
            ("权利取得方式", "原始取得"),
            ("权利范围", "全部权利"),
            ("共同申请原则", "一名实际开发成员一行；不自动纳入非开发参赛成员或指导教师"),
            ("申请代表", "【申请代表姓名】；电话【申请代表电话】；邮箱【申请代表邮箱】"),
        ],
    )
    doc.add_heading("三、共同著作权人录入表", level=1)
    rows = []
    for i in range(1, 9):
        rows.append(
            (
                str(i),
                f"【共同著作权人{i}姓名】",
                f"【证件类型/号码{i}】",
                f"【地址/电话/邮箱{i}】",
                f"【实际开发贡献{i}】",
            )
        )
    add_table(doc, ["序号", "姓名", "证件信息", "联系信息", "实际开发贡献"], rows, 7.5)
    doc.add_paragraph("说明：不足八人时删除空行；超过八人时按相同字段扩展。所有成员必须属于实际开发人员。")
    doc.add_heading("四、软件功能与用途", level=1)
    doc.add_paragraph(
        "本软件面向基于电路仿真数据的工程评估与参数迭代场景。系统以工作区、项目和设计版本组织仿真输入，"
        "对波形 CSV、参数 YAML、网表及图像工件进行预览与校验；调用约束评估内核形成问题、指标和证据索引；"
        "管理机器生成或规则生成的候选参数；在人工审批后导出可复现仿真任务；接收外部仿真结果回填并创建新版本；"
        "最终对基线与结果版本进行比较，输出可追溯报告。"
    )
    add_boundary(doc)
    doc.add_heading("五、技术特点", level=1)
    add_bullets(
        doc,
        [
            "采用 Python/FastAPI/SQLAlchemy 构建领域服务与 REST API，采用 React/TypeScript 构建浏览器界面。",
            "以工作区、项目、设计版本、分析运行、实验、候选、仿真任务和比较结果构成可追溯状态模型。",
            "结构化数据库保存业务状态，内容寻址工件库保存 CSV、JSON、YAML、网表、图像和报告。",
            "输入预览与正式提交分离；状态转换受领域规则约束；错误采用稳定错误码与详细信息返回。",
            "候选审批只授权导出，不自动运行外部仿真；回填结果经过清单和结构校验后才形成新版本。",
            "证据边界贯穿 API、界面、报告和材料，防止把仿真证据写成实测或流片结论。",
        ],
    )
    doc.add_heading("六、开发与运行环境", level=1)
    add_kv_table(
        doc,
        [
            ("开发硬件环境", "通用 x86-64 个人计算机；建议内存8 GB以上；可用磁盘空间2 GB以上"),
            ("运行硬件环境", "通用 x86-64 个人计算机或服务器；建议内存4 GB以上"),
            ("开发操作系统", "Windows 10/11；兼容常见 Linux 开发环境"),
            ("运行操作系统", "Windows 10/11 或主流 Linux 发行版"),
            ("后端环境", "Python 3.10及以上；FastAPI；Uvicorn；SQLAlchemy；SQLite"),
            ("前端环境", "Node.js；React；TypeScript；Vite；现代浏览器"),
            ("编程语言", "Python、TypeScript/TSX、CSS"),
            ("源程序量", f"{source_stats['files']}个自研源文件，合计{source_stats['lines']}行"),
        ],
    )
    doc.add_heading("七、版本与发表说明", level=1)
    doc.add_paragraph(
        f"申报版本以 Git 提交 {BASELINE_COMMIT} 为功能基线，在分支 {BRANCH} "
        f"仅对 Python 包、前端包及锁文件的版本元数据统一为 {PACKAGE_VERSION}，未改变产品 API、CLI、数据库模型"
        f"或前端业务行为。首次发表口径为已于{FIRST_PUBLICATION_DATE}在 GitHub 公开仓库发表；正式提交前由团队保存和核对"
        "可证明公开日期、仓库归属与作品内容的证据。"
    )
    doc.add_heading("八、官方系统录入核对表", level=1)
    add_table(
        doc,
        ["核对项", "本底稿值", "录入后复核"],
        [
            ("全称/简称/版本", f"{FULL_NAME} / {SHORT_NAME} / {VERSION}", "□"),
            ("完成与发表日期", f"{COMPLETION_DATE} / {FIRST_PUBLICATION_DATE}", "□"),
            ("开发及权利", "合作开发 / 原始取得 / 全部权利", "□"),
            ("著作权人", "仅实际开发成员，逐人一行", "□"),
            ("软件功能与技术特点", "不得删改证据边界", "□"),
            ("源码量", f"{source_stats['files']}文件 / {source_stats['lines']}行", "□"),
        ],
        8.5,
    )
    doc.add_paragraph(f"官方办理规范：{OFFICIAL_FORM_URL}")
    path = output / "01_软件著作权登记申请表填报底稿.docx"
    save_doc(doc, path)
    return path


def build_design_spec(output: Path, source_stats: dict) -> Path:
    doc = Document()
    configure_document(doc)
    add_cover(doc, "业务理解与软件设计说明书", "从业务闭环、系统架构到证据链的完整说明", "设计说明书")
    doc.add_heading("文档控制", level=1)
    add_kv_table(
        doc,
        [
            ("文档版本", "1.0"),
            ("对应软件", f"{FULL_NAME} {VERSION}"),
            ("功能基线", BASELINE_COMMIT),
            ("适用读者", "共同开发成员、登记材料编制人员、软件评审人员、工程操作员"),
            ("保密级别", "申请材料；不含个人身份信息、密钥与私有数据"),
        ],
    )
    doc.add_heading("1. 业务背景与目标", level=1)
    doc.add_paragraph(
        "电路参数迭代通常跨越波形文件、参数表、约束定义、仿真器输出、人工判断和多轮版本。"
        "如果这些信息只分散在目录和聊天记录中，容易出现输入版本不明、候选来源不明、重仿真遗漏以及结论越过证据边界。"
        "芯智调参把这些环节组织为可追溯的软件流程，使每一个分析结论都能回到输入、规则、工件和状态。"
    )
    doc.add_paragraph(
        "V1.0 的目标不是替代电路工程师或仿真器，而是提供仿真数据管理、约束分析、候选参数推荐、人工审批、"
        "仿真任务导出、结果回填和版本比较的一体化工作台。外部仿真仍由工程操作员控制。"
    )
    add_boundary(doc)
    doc.add_heading("2. 术语与边界", level=1)
    add_table(
        doc,
        ["术语", "定义", "不代表"],
        [
            ("设计版本", "一组可追溯的参数、网表和输入引用", "已流片芯片版本"),
            ("分析运行", "针对特定设计版本执行的约束与指标计算", "实验室实测"),
            ("候选参数", "供下一轮仿真选择的参数变化建议", "已证实改进"),
            ("仿真任务", "可导出、可回填、受状态机控制的工程任务", "自动运行所有外部工具"),
            ("比较结论", "基于已回填仿真证据的结构化差异", "保证性能提升"),
            ("真实仿真 CSV", "由外部仿真过程形成并按契约导入的 CSV", "真实芯片测量"),
        ],
        8.2,
    )
    doc.add_heading("3. 用户角色与职责", level=1)
    add_table(
        doc,
        ["角色", "主要职责", "权限边界"],
        [
            ("工程操作员", "创建项目、上传输入、查看分析、审批候选、导出任务、回填结果", "对外部仿真和工程判断负责"),
            ("评审人员", "查看证据链、问题、比较和报告", "只读评审，不把候选改写为验证结论"),
            ("开发维护人员", "维护代码、接口、规则、测试和部署", "不得绕过状态机或隐藏证据来源"),
            ("申请材料编制人员", "依据固定提交汇编源码和文档", "不得生成、改写或补造申报源码"),
        ],
        8.2,
    )
    doc.add_heading("4. 核心业务流程", level=1)
    add_numbered(
        doc,
        [
            "创建工作区和项目，选择电路配置与规格修订。",
            "创建基线设计版本，为参数、网表和输入工件建立稳定引用。",
            "上传仿真波形 CSV、参数 YAML；网表与图像作为可选证据工件。",
            "执行输入预览，检查文件类型、字段、清单和证据边界。",
            "正式创建分析运行，生成指标、约束问题、证据索引和运行清单。",
            "创建实验并生成候选参数；工程人员审阅后批准、拒绝或保留候选。",
            "将批准候选组成手工仿真任务，导出可复现批次。",
            "工程人员在受控外部仿真器中运行批次，系统不自动宣称仿真成功。",
            "导入结果 CSV，先预览校验，再提交形成结果设计版本与新分析运行。",
            "比较基线与结果版本，输出差异、约束变化和中性/改进/退化判定。",
        ],
    )
    doc.add_heading("5. 总体架构", level=1)
    doc.add_paragraph(
        "系统采用浏览器前端、REST API、领域服务、持久化仓储与内容工件库的分层架构。React/TypeScript 前端"
        "负责交互和状态呈现；FastAPI 路由负责协议适配与错误映射；领域服务负责规则和状态转换；SQLAlchemy 仓储"
        "保存业务对象；LocalArtifactStore 保存大体量或可复现实物。"
    )
    add_table(
        doc,
        ["层次", "组成", "职责"],
        [
            ("表示层", "React、TypeScript、React Router、Recharts", "页面路由、表单、状态、证据与比较呈现"),
            ("接口层", "FastAPI Product API", "HTTP 契约、请求校验、响应模型、错误处理"),
            ("领域层", "Project/Input/Analysis/Experiment/Simulation/Comparison Service", "业务规则与状态机"),
            ("评估内核", "goa_eval、PIA、CA、LLSO、profiles", "波形读取、指标、约束和候选逻辑"),
            ("持久化层", "SQLAlchemy、SQLite、LocalArtifactStore", "结构化状态与不可变工件引用"),
            ("工程接口", "CLI、导出包、CSV/YAML/JSON", "批处理、复现与外部仿真衔接"),
        ],
        8,
    )
    doc.add_heading("6. 模块职责", level=1)
    group_counts = Counter(item["source_group"] for item in source_stats["entries"])
    group_lines = Counter()
    for item in source_stats["entries"]:
        group_lines[item["source_group"]] += item["line_count"]
    add_table(
        doc,
        ["模块组", "文件数", "代码行", "主要职责"],
        [
            (
                group,
                str(group_counts[group]),
                str(group_lines[group]),
                {
                    "产品 API": "REST 路由、请求响应模型、异常映射和依赖容器",
                    "产品服务": "项目、输入、分析、实验、任务、比较、仓储和工件",
                    "只读仪表盘 API": "演示包与报告读取接口",
                    "上传分析服务": "文件上传、预览和分析后端",
                    "PIA 与候选生成": "候选生成、训练、评估和适配",
                    "电路分析": "电路分析与指标处理",
                    "LLSO 优化": "搜索与优化辅助",
                    "多智能体协作": "任务协调与结构化分析",
                    "评估内核与命令行": "波形读取、约束、评分、报告和 CLI",
                    "前端界面": "路由、页面、组件、API 客户端与样式",
                    "工程脚本": "复现、演示、数据与维护脚本",
                }.get(group, "工程实现"),
            )
            for group in sorted(group_counts, key=lambda g: min(i["_rank"] for i in source_stats["entries"] if i["source_group"] == g))
        ],
        8,
    )
    doc.add_heading("7. 数据模型", level=1)
    add_table(
        doc,
        ["实体", "关键标识", "关键关系/状态"],
        [
            ("Workspace", "workspace_id", "拥有多个 Project"),
            ("Project", "project_id", "隶属 Workspace；引用 profile/spec"),
            ("DesignVersion", "design_version_id", "父版本、来源候选、参数/网表引用"),
            ("AnalysisRun", "analysis_run_id", "隶属版本；pending/running/completed/failed"),
            ("Experiment", "experiment_id", "绑定基线版本与生成策略"),
            ("Candidate", "candidate_id", "proposed/approved/rejected/exported/evaluated"),
            ("SimulationJob", "job_id", "draft/queued/exported/completed/failed"),
            ("Comparison", "comparison_id", "基线版本、结果版本、差异与判定"),
            ("ArtifactRef", "artifact:// URI + SHA-256", "文件路径不直接暴露给业务层"),
        ],
        8.2,
    )
    doc.add_heading("8. 数据流与工件链", level=1)
    add_numbered(
        doc,
        [
            "输入文件写入工件库并计算 SHA-256；数据库只保存不可变引用与业务元数据。",
            "输入清单记录波形、参数、网表、图片、配置修订和边界字段。",
            "分析服务从清单读取输入，生成运行清单、指标、问题、证据索引与报告。",
            "候选记录保存参数差异、策略、分数和 must_resimulate 状态。",
            "仿真任务导出包保存候选集合、基线信息、预期结果契约和校验摘要。",
            "回填服务验证任务清单与 CSV 结构，提交后创建结果版本并触发分析。",
            "比较服务仅使用已完成分析运行，生成指标差异与约束变化。",
        ],
    )
    doc.add_heading("9. API 设计", level=1)
    api_rows = [
        ("GET/POST", "/api/v1/workspaces", "列出或创建工作区"),
        ("GET/POST", "/api/v1/workspaces/{workspace_id}/projects", "工作区项目"),
        ("GET", "/api/v1/projects/{project_id}/overview", "项目版本与分析概览"),
        ("GET/POST", "/api/v1/projects/{project_id}/design-versions", "设计版本"),
        ("POST", "/api/v1/design-versions/{version_id}/inputs/preview", "输入预览"),
        ("GET/POST", "/api/v1/design-versions/{version_id}/analysis-runs", "分析运行"),
        ("GET", "/api/v1/analysis-runs/{run_id}/issues|evidence|bundle", "问题、证据与工件包"),
        ("GET/POST", "/api/v1/projects/{project_id}/experiments", "实验"),
        ("POST", "/api/v1/experiments/{experiment_id}/candidates:generate", "候选生成"),
        ("POST", "/api/v1/candidates/{candidate_id}:approve|reject|confirm", "候选状态转换"),
        ("POST", "/api/v1/simulation-jobs", "创建仿真任务"),
        ("POST", "/api/v1/simulation-jobs/{job_id}:export", "导出任务"),
        ("POST", "/api/v1/simulation-jobs/{job_id}/imports:preview|commit", "结果预览与提交"),
        ("POST/GET", "/api/v1/comparisons", "创建与读取版本比较"),
    ]
    add_table(doc, ["方法", "路径", "职责"], api_rows, 7.5)
    doc.add_heading("10. 状态机设计", level=1)
    add_table(
        doc,
        ["对象", "允许转换", "阻止条件"],
        [
            ("AnalysisRun", "pending→running→completed/failed", "输入不完整或运行异常"),
            ("Candidate", "proposed→approved/rejected；approved→exported；exported→evaluated", "未审批不得导出；无重仿真不得评价"),
            ("SimulationJob", "draft→queued/exported→completed/failed", "候选未批准、清单不一致或回填校验失败"),
            ("Comparison", "created→evaluated", "基线或结果分析未完成"),
        ],
        8,
    )
    doc.add_paragraph(
        "候选、仿真任务和比较对象共同构成“建议—审批—外部仿真—回填—评价”闭环。任何状态都不能把"
        " must_resimulate = true 静默改为已验证。"
    )
    doc.add_heading("11. 异常处理", level=1)
    add_table(
        doc,
        ["类别", "处理方式", "用户可恢复动作"],
        [
            ("请求校验", "FastAPI/Pydantic 返回字段级错误", "修正字段后重试"),
            ("业务冲突", "稳定错误码与当前状态", "刷新对象并按允许状态执行"),
            ("输入格式", "预览阶段报告缺列、类型和清单错误", "修正 CSV/YAML 后重新预览"),
            ("工件缺失", "拒绝读取并保留 URI/摘要用于追踪", "恢复工件或重新上传"),
            ("仿真失败", "任务进入 failed；不创建结果版本", "修正外部仿真后重试或新建任务"),
            ("比较不足", "返回不可判定或中性结果", "补充完成的结果运行"),
            ("未知异常", "统一异常处理并记录服务端诊断", "由维护人员定位；不向页面泄露密钥"),
        ],
        8,
    )
    doc.add_heading("12. 证据链设计", level=1)
    add_boundary(doc)
    doc.add_paragraph(
        "每次运行通过输入清单、配置修订、工件 URI、SHA-256、状态时间戳、问题记录、证据索引和报告形成链路。"
        "界面展示的数据来源于这些持久化对象，不依赖人工复制的成功描述。"
    )
    doc.add_heading("13. 安全与隐私", level=1)
    add_bullets(
        doc,
        [
            "文件类型、路径和工件 URI 受服务层约束，避免任意路径读取。",
            "源码材料排除环境变量、私有数据、依赖目录、锁文件、测试数据和生成输出。",
            "操作手册截图使用 test_only 演示库，不含真实身份、密钥和个人联系信息。",
            "系统输出不应包含本地绝对路径；异常信息在界面层进行结构化呈现。",
            "部署时应通过反向代理、访问控制、最小权限文件目录和备份策略补充生产安全措施。",
        ],
    )
    doc.add_heading("14. 部署与运行环境", level=1)
    add_kv_table(
        doc,
        [
            ("后端", "Python 3.10+；Uvicorn 启动 FastAPI；SQLite 可用于单机部署"),
            ("前端", "Node.js 构建静态资源；现代 Chromium/Edge/Firefox 浏览器"),
            ("存储", "结构化数据库 + 本地工件目录；生产环境应配置持久卷和备份"),
            ("网络", "前端与 API 同源或经允许的代理；外部仿真器可保持离线"),
            ("可观测性", "HTTP 状态、稳定错误码、运行状态、时间戳与工件摘要"),
        ],
    )
    doc.add_heading("15. 测试与质量保证", level=1)
    doc.add_paragraph(
        "申报分支重新执行后端 pytest、前端 Vitest 与生产构建，并运行确定性产品演示。"
        "演示数据库只用于验证工作流和界面，不被描述为真实实验验证。固定分页材料另行执行页数、行数、页眉、页码、"
        "文本一致性、DOCX 渲染与 PDF 再渲染检查。"
    )
    doc.add_heading("16. 版本管理与变更控制", level=1)
    doc.add_paragraph(
        f"V1.0 以 {BASELINE_COMMIT} 为功能基线。申报分支仅统一 pyproject.toml、frontend/package.json "
        "与 frontend/package-lock.json 的版本元数据，并新增材料生成/校验工具。产品 API、CLI、数据库模型和前端业务行为不变。"
    )
    doc.add_heading("17. 限制、风险与后续演进", level=1)
    add_bullets(
        doc,
        [
            "当前工程有效性为 simulation_only，不能外推为芯片级实测结论。",
            "候选质量受输入数据、约束配置、搜索策略和外部仿真质量影响。",
            "SQLite 与本地工件库适合单机或演示环境，大规模协作需评估数据库和对象存储。",
            "外部仿真器、PDK 和许可证由使用者自行配置，系统不包含其著作权或使用授权。",
            "后续可增加身份认证、审计日志、任务队列、更多电路配置和更严格的结果签名。",
        ],
    )
    doc.add_heading("附录A：源程序规模", level=1)
    doc.add_paragraph(f"确定性清单共 {source_stats['files']} 个自研源文件、{source_stats['lines']} 行。")
    add_table(
        doc,
        ["语言", "文件数", "行数"],
        [
            (lang, str(sum(1 for e in source_stats["entries"] if e["language"] == lang)),
             str(sum(e["line_count"] for e in source_stats["entries"] if e["language"] == lang)))
            for lang in sorted({e["language"] for e in source_stats["entries"]})
        ],
        8.5,
    )
    doc.add_heading("附录B：规范参考", level=1)
    doc.add_paragraph(f"办理规范：{OFFICIAL_FORM_URL}")
    doc.add_paragraph(f"计算机软件著作权登记办法：{REGULATION_URL}")
    path = output / "02_业务理解与软件设计说明书.docx"
    save_doc(doc, path)
    return path


def build_manual(output: Path, screenshot_dir: Path) -> Path:
    doc = Document()
    configure_document(doc)
    add_cover(doc, "软件操作手册", "安装、运行与仿真闭环操作说明", "用户操作手册")
    add_notice(
        doc,
        "截图来源",
        "本手册界面截图均由 Playwright 在 1440×900 统一视口下操作真实运行系统获得。"
        "所用数据库与文件为确定性测试夹具 test_only，仅验证软件工作流，不代表真实实验、芯片实测或流片验证。",
    )
    doc.add_heading("1. 软件概述", level=1)
    doc.add_paragraph(
        "芯智调参用于组织电路仿真输入、执行约束分析、管理候选参数、导出人工仿真任务、"
        "回填结果并比较设计版本。系统由 Python/FastAPI 后端和 React 浏览器前端组成。"
    )
    add_boundary(doc)
    doc.add_heading("2. 安装前准备", level=1)
    add_bullets(
        doc,
        [
            "Windows 10/11 或主流 Linux；Python 3.10及以上；Node.js 与 npm。",
            "从固定提交检出源码，确认工作目录不包含私人数据和密钥。",
            "外部仿真器及其 PDK/许可证不随本软件提供，由工程操作员单独配置。",
            "建议使用虚拟环境安装 Python 依赖；前端使用 npm ci 安装锁定依赖。",
        ],
    )
    doc.add_heading("3. 安装与启动", level=1)
    for command in [
        "python -m venv .venv",
        ".venv\\Scripts\\activate",
        "python -m pip install -e .",
        "cd frontend && npm ci",
        "python -m uvicorn goa_eval.product_api.app:app --host 127.0.0.1 --port 8001",
        "cd frontend && npm run dev -- --host 127.0.0.1 --port 5173",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        r = p.add_run(command)
        set_run_font(r, east_asia="宋体", latin="Courier New", size=9)
    doc.add_paragraph("浏览器访问 http://127.0.0.1:5173。生产环境应使用前端构建产物并配置同源反向代理。")

    screenshot_specs = [
        ("09_create_project.png", "4. 创建项目", [
            "进入 Projects 页面并选择 New project。",
            "填写项目名称，选择电路配置与规格修订。",
            "核对证据边界后提交。项目创建成功会返回项目标识。",
        ]),
        ("08_upload_workspace.png", "5. 上传仿真结果与创建分析上下文", [
            "进入 Upload analysis，先选择工作区和已有项目，或创建新项目。",
            "填写设计版本名称；上传 waveform CSV 与 parameter YAML。",
            "可选上传网表和图像证据。先执行预览，修正字段错误后再提交分析。",
        ]),
        ("03_design_version.png", "6. 设计版本与输入预览", [
            "打开设计版本页面，确认版本标签与版本标识。",
            "选择波形、参数、网表和图片；Advanced settings 用于受控高级参数。",
            "点击 Preview input。预览只校验输入，不形成已验证工程结论。",
        ]),
        ("04_analysis_result.png", "7. 查看分析结果", [
            "分析运行完成后查看运行状态、约束问题与证据索引。",
            "问题列表给出规则、严重度和已知性；证据项可回溯到工件。",
            "任何建议继续保持 must_resimulate = true。",
        ]),
        ("05_candidate_approval.png", "8. 候选审批", [
            "进入优化实验页，查看策略、基线版本、参数变化、分数和状态。",
            "审批只授权把候选加入仿真导出批次，不自动执行仿真。",
            "不合适的候选可拒绝；待补充信息的候选保持 proposed。",
        ]),
        ("06_simulation_job.png", "9. 导出仿真任务与结果回填", [
            "从已审批候选创建 Manual simulation job，并导出可复现批次。",
            "在外部仿真器中运行批次；系统不控制外部仿真器。",
            "返回后选择 Simulation result CSV，先 Preview import，再 Commit import。",
            "提交成功才创建结果设计版本并进入后续分析。",
        ]),
        ("07_version_comparison.png", "10. 版本比较", [
            "选择基线版本与结果版本创建比较。",
            "查看总体判定、指标差异和约束变化。",
            "示例 test_only 数据的判定为 No material change，不宣称优化成功。",
        ]),
        ("01_public_demo.png", "11. 报告与演示仪表盘", [
            "Public demo 汇总运行状态、评分、硬约束和证据边界。",
            "下方可查看候选排序、重跑状态、图像包和报告交付。",
            "演示页明确说明仿真证据不代表物理验证、芯片验证或流片验证。",
        ]),
    ]
    for filename, title, steps in screenshot_specs:
        doc.add_heading(title, level=1)
        add_numbered(doc, steps)
        path = screenshot_dir / filename
        if not path.is_file():
            raise FileNotFoundError(path)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(6.4))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"图 {filename[:2]}  {title.split('. ', 1)[-1]}（test_only 真实运行截图）")
        set_run_font(r, east_asia="宋体", size=8.5, color=GRAY)

    doc.add_heading("12. 常见异常与处理", level=1)
    add_table(
        doc,
        ["现象", "原因", "处理"],
        [
            ("资源加载失败", "API 地址、代理、服务或对象标识不正确", "检查 API 进程、前端代理和对象标识"),
            ("预览失败", "CSV/YAML 缺列、格式或类型不符", "按错误详情修正后重新预览"),
            ("候选不可导出", "候选未审批或状态冲突", "刷新状态并先完成审批"),
            ("回填不可提交", "任务清单、候选或结果结构不一致", "使用对应导出包重新生成结果"),
            ("比较不可创建", "基线或结果分析未完成", "等待分析完成或检查失败原因"),
            ("无改进结论", "指标差异不足或约束未改善", "保留中性结论，不人为改写结果"),
        ],
        8,
    )
    doc.add_heading("13. 数据与安全注意事项", level=1)
    add_bullets(
        doc,
        [
            "上传前移除 CSV/YAML 中的个人信息、密钥和不必要的本地路径。",
            "不要把 test_only 夹具写成真实实验数据。",
            "备份数据库和工件目录时保持二者一致。",
            "导出任务应在受控目录运行，外部仿真结果应保留原始文件和时间记录。",
            "报告中的候选建议必须与证据边界同时解释。",
        ],
    )
    doc.add_heading("14. 退出与停止", level=1)
    doc.add_paragraph("在前端和后端终端分别按 Ctrl+C 停止服务。确认无写入任务后再备份或迁移数据库与工件目录。")
    path = output / "03_软件操作手册.docx"
    save_doc(doc, path)
    return path


def build_joint_agreement(output: Path) -> Path:
    doc = Document()
    configure_document(doc)
    add_cover(doc, "合作开发与著作权归属确认书", "实际开发成员共同申请模板", "需逐人签署的权属文件")
    doc.add_heading("一、确认事项", level=1)
    doc.add_paragraph(
        f"各方确认共同开发软件《{FULL_NAME}》（简称“{SHORT_NAME}”，版本号{VERSION}）。"
        "本确认书用于明确实际开发成员之间的合作开发事实、贡献、权属和登记申请安排。"
    )
    doc.add_heading("二、共同开发成员与贡献", level=1)
    rows = []
    for i in range(1, 9):
        rows.append((str(i), f"【成员{i}姓名】", f"【成员{i}实际开发贡献】", f"【成员{i}证据索引】", f"【成员{i}签字】"))
    add_table(doc, ["序号", "成员姓名", "实际开发贡献", "提交/任务/设计证据", "签字"], rows, 7.3)
    doc.add_paragraph("仅填写对源代码、界面或技术文档作出实际创作贡献的成员；不足或超过八人时据实调整。")
    doc.add_heading("三、著作权归属", level=1)
    add_numbered(
        doc,
        [
            "各方确认本软件由表列实际开发成员合作开发，著作权由实际开发成员共同享有。",
            "登记申请采用原始取得、全部权利口径；各方不得就本确认书所列共同作品单独作出与共同权属冲突的声明。",
            "第三方开源组件、运行环境、外部仿真器、PDK及其文档不属于各方共同原创权利范围。",
            "如任何成员的贡献受劳动关系、单位任务、委托合同、既有协议或第三方权利约束，该成员应在签署前书面披露并完成必要授权。",
        ],
    )
    doc.add_heading("四、申请代表与办理授权", level=1)
    doc.add_paragraph(
        "各方推选【申请代表姓名】作为登记办理联系人，负责材料汇总、系统录入、补正沟通和证书接收。"
        "申请代表不得擅自变更共同著作权人、权利范围、软件名称、版本或处分安排。"
    )
    doc.add_heading("五、收益与处分约定", level=1)
    add_numbered(
        doc,
        [
            "软件许可、转让、质押、开源、重大商业合作等处分事项，应由共同著作权人依法律规定及另行书面约定共同决定。",
            "与软件有关的收益分配、成本承担和维权费用由各方另行书面约定；未约定事项依法协商处理。",
            "任何一方拟向第三方披露未公开源码、私有数据或密钥，应先取得必要授权并遵守保密义务。",
        ],
    )
    doc.add_heading("六、陈述与保证", level=1)
    add_bullets(
        doc,
        [
            "所填贡献真实、可核验，不把非开发参赛成员或指导教师自动列为著作权人。",
            "提交的原创代码、界面和文档未故意侵犯他人著作权。",
            "申报材料不把仿真结果宣称为物理实测、芯片验证或流片证明。",
            "各方已阅读申请表底稿、源程序鉴别材料和第三方权利边界说明。",
        ],
    )
    doc.add_heading("七、争议解决与生效", level=1)
    doc.add_paragraph(
        "因本确认书产生的争议，由各方先行协商；协商不成时，依法向有管辖权的人民法院处理。"
        "本确认书自全体实际开发成员签字之日起生效，可按成员人数制作同等效力文本。"
    )
    doc.add_heading("八、签署页", level=1)
    for i in range(1, 9):
        doc.add_paragraph(f"共同开发成员{i}：【姓名】        签字：【签字】        日期：【签字日期】")
        doc.add_paragraph(f"证件号码：【证件号码】          联系电话：【联系电话】")
    doc.add_paragraph("申请代表：【姓名】              签字：【签字】        日期：【签字日期】")
    doc.add_paragraph(
        "提示：本模板不能替代对劳动关系、单位任务、委托开发、既有合同等特殊权属情形的专业法律审查。"
    )
    path = output / "05_合作开发与著作权归属确认书.docx"
    save_doc(doc, path)
    return path


def build_third_party(output: Path) -> Path:
    doc = Document()
    configure_document(doc)
    add_cover(doc, "第三方组件与权利边界说明", "原创范围、依赖用途与申报边界", "内部支撑/备查")
    doc.add_heading("一、目的与声明", level=1)
    doc.add_paragraph(
        "本说明用于区分团队原创软件与第三方运行库、开发工具及外部工程资源。"
        "本次登记申请仅主张固定提交中由团队原创的业务代码、界面组织与技术文档，不主张第三方组件本身的著作权。"
    )
    components = [
        ("Python", "后端语言与运行时", "Python Software Foundation License 等；以所用版本随附文本为准"),
        ("FastAPI", "REST API 框架", "MIT；以安装包许可证为准"),
        ("Uvicorn", "ASGI 服务运行", "BSD 类；以安装包许可证为准"),
        ("Pydantic", "请求、响应与配置校验", "MIT；以安装包许可证为准"),
        ("SQLAlchemy", "数据库映射与事务", "MIT；以安装包许可证为准"),
        ("NumPy", "数值数组与计算", "BSD 类；以安装包许可证为准"),
        ("Pandas", "CSV 与表格数据处理", "BSD 类；以安装包许可证为准"),
        ("Matplotlib", "工程图表生成", "Matplotlib License；以安装包文本为准"),
        ("React", "浏览器组件框架", "MIT；以 npm 包许可证为准"),
        ("React Router", "前端页面路由", "MIT；以 npm 包许可证为准"),
        ("Recharts", "前端图表", "MIT；以 npm 包许可证为准"),
        ("Lucide React", "界面图标", "ISC；以 npm 包许可证为准"),
        ("Vite", "前端开发与构建", "MIT；以 npm 包许可证为准"),
        ("Tailwind CSS", "样式构建工具", "MIT；以 npm 包许可证为准"),
        ("Vitest", "前端自动化测试", "MIT；仅开发测试使用"),
        ("Playwright", "浏览器验收与截图", "Apache-2.0；仅开发测试与材料采集使用"),
        ("SQLite", "单机结构化数据存储", "公共领域；以官方声明为准"),
    ]
    doc.add_heading("二、主要第三方组件清单", level=1)
    add_table(doc, ["组件", "用途", "许可证提示"], components, 7.5)
    doc.add_paragraph(
        "许可证名称用于材料边界说明，不替代对最终分发包的逐版本许可证审计。正式发布或商业分发前应根据锁定版本、"
        "安装包和上游 NOTICE/LICENSE 文件再次核对。"
    )
    doc.add_heading("三、不纳入权利主张的内容", level=1)
    add_bullets(
        doc,
        [
            "Python、FastAPI、React、SQLAlchemy 等第三方组件及其源代码和文档。",
            "Node.js、npm、浏览器、操作系统、数据库引擎等运行环境。",
            "外部 SPICE/EDA 仿真器、PDK、模型库、许可证和厂商文档。",
            "开源字体、图标、测试框架及材料生成工具中的第三方库。",
            "用户上传的仿真数据、网表和图像中属于用户或第三方的内容。",
        ],
    )
    doc.add_heading("四、团队原创范围", level=1)
    add_bullets(
        doc,
        [
            "工作区、项目、版本、分析、实验、候选、仿真任务和比较等业务模型与流程实现。",
            "输入预览、证据索引、状态机、工件引用、错误处理和边界呈现等原创业务逻辑。",
            "CircuitPilot 产品 API 的路由组织、领域服务、仓储适配和报告逻辑。",
            "芯智调参前端页面、交互流程、信息架构和团队原创样式。",
            "本次提交的业务理解、设计说明、操作说明和权利边界文档。",
        ],
    )
    doc.add_heading("五、源码鉴别材料排除规则", level=1)
    add_bullets(
        doc,
        [
            "排除 node_modules、虚拟环境、缓存、构建产物和第三方依赖源码。",
            "排除测试目录与测试文件、锁文件、输出数据、环境变量文件和私有数据。",
            "排除自动生成样板和本次材料生成工具。",
            "只使用 Git 跟踪的 Python、TypeScript/TSX、CSS 自研文件，并记录 SHA-256。",
        ],
    )
    doc.add_heading("六、证据边界", level=1)
    add_boundary(doc)
    doc.add_paragraph(
        "第三方仿真器或 PDK 的存在不改变本软件的证据口径。系统只对导入工件和软件流程负责，"
        "不对外部组件的正确性、许可证或工程结果作保证。"
    )
    path = output / "06_第三方组件与权利边界说明.docx"
    save_doc(doc, path)
    return path


def _new_template_doc() -> Document:
    doc = Document()
    doc._body.clear_content()
    configure_document(doc, compact=True)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(4)
    return doc


def _clear_container(container) -> None:
    for child in list(container._element):
        container._element.remove(child)


def _begin_template_page(
    doc: Document,
    physical_page: int,
    total_pages: int,
    kind: str,
    *,
    cover: bool = False,
    body_page: int | None = None,
    body_total: int | None = None,
    profile: str = "standard",
):
    section = doc.sections[0] if physical_page == 1 else doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    if profile == "application":
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.0)
    else:
        section.top_margin = Cm(1.75)
        section.bottom_margin = Cm(1.65)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    _clear_container(section.header)
    _clear_container(section.footer)
    if cover:
        return section

    shown_page = body_page if body_page is not None else physical_page
    shown_total = body_total if body_total is not None else total_pages
    header = section.header
    table = header.add_table(rows=1, cols=3, width=Cm(17.0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    values = (f"{FULL_NAME} {VERSION}", kind, f"第{shown_page}页，共{shown_total}页")
    widths = (9.4, 3.4, 4.2)
    aligns = (WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT)
    for index, value in enumerate(values):
        cell = table.cell(0, index)
        cell.width = Cm(widths[index])
        paragraph = cell.paragraphs[0]
        paragraph.alignment = aligns[index]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(value)
        set_run_font(run, east_asia="宋体", size=7.3, bold=(index == 1), color=BLACK)
        if index == 0:
            run._element.get_or_add_rPr().append(_char_scale(78))
    footer = section.footer.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run(f"{shown_page}/{shown_total}")
    set_run_font(run, east_asia="宋体", size=8, color=BLACK)
    return section


def _add_template_cover(doc: Document, title: str, subtitle: str, material_use: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(105)
    r = p.add_run(FULL_NAME)
    set_run_font(r, east_asia="黑体", size=19, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    r = p.add_run(title)
    set_run_font(r, east_asia="黑体", size=24, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run(subtitle)
    set_run_font(r, east_asia="宋体", size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(115)
    for line in (
        f"软件简称：{SHORT_NAME}",
        f"版本号：{VERSION}",
        f"材料用途：{material_use}",
        "申请口径：实际开发成员共同申请",
        "编制日期：2026年07月",
    ):
        r = p.add_run(line + "\n")
        set_run_font(r, east_asia="宋体", size=11)


def _add_page_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title)
    set_run_font(r, east_asia="黑体", size=16, bold=True)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(subtitle)
        set_run_font(r, east_asia="宋体", size=9, color=GRAY)


def _add_body_paragraphs(doc: Document, paragraphs: Sequence[str]) -> None:
    for text in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.45
        r = p.add_run(text)
        set_run_font(r, east_asia="宋体", size=10.5)


def build_template_catalog(output: Path, source_stats: dict) -> Path:
    doc = _new_template_doc()
    pages = [
        (
            "材料总目录与用途",
            [
                ("正式提交候选", "03操作手册全文（文档鉴别材料）、04源程序前后各30页；按受理系统要求提交。"),
                ("官方系统生成件", "将01底稿内容录入中国版权保护中心系统后生成正式申请表并签章；自制底稿不能替代。"),
                ("权属证明", "05合作合同及05A共同权属确认书，经实际开发成员逐人核实并签署后使用。"),
                ("内部支撑件", "00目录、02设计说明书、06第三方边界、清单与internal_support用于核对和备查。"),
            ],
        ),
        (
            "正式材料与支撑材料清单",
            [
                ("01", "申请表填报底稿；六页；非官方申请表；不得制作二维码或官方签章。"),
                ("03", "18页完整操作手册；不足60页按全文提交；封面不编号、正文17页连续编号。"),
                ("04", "60页源程序；每页50行；显示行号1—3000；前后片段分别连续。"),
                ("05/05A", "合作开发合同与共同权属确认书；补齐身份、贡献、代表、地点、日期和签章。"),
                ("02/06/清单", "业务与设计、第三方边界、哈希和追溯数据；默认内部留存。"),
            ],
        ),
        (
            "提交前检查清单",
            [
                ("申请人", "【共同申请人姓名】、【证件号码】、【地址】、【电话】逐人填写并与证件一致。"),
                ("发表信息", f"【首次发表城市】；完成日与首次发表日均为{COMPLETION_DATE}；方式为GitHub公开仓库。"),
                ("签署", "合同编号、签订地点、申请代表、贡献说明、签字/盖章及日期必须由团队确认。"),
                ("一致性", f"全称、简称、{VERSION}、合作开发、原始取得、全部权利保持一致。"),
                ("边界", "不得宣称流片、物理实测或效果保证；候选参数必须经外部仿真复核。"),
                ("源码统计", f"本次纳入{source_stats['files']}个自研文件、{source_stats['lines']}行原始源码；正式鉴别材料抽取3000行。"),
            ],
        ),
    ]
    for index, (title, rows) in enumerate(pages, 1):
        _begin_template_page(doc, index, 3, "材料总目录与提交检查清单")
        _add_page_title(doc, title, "模板对齐版申报包内部总控")
        add_kv_table(doc, rows, widths=(4.0, 12.0))
        if index in (1, 3):
            add_boundary(doc)
    path = output / "00_材料总目录与提交检查清单.docx"
    save_doc(doc, path)
    return path


def build_template_application(output: Path, source_stats: dict) -> Path:
    doc = _new_template_doc()
    page_rows = [
        ("一、软件基本信息", [("软件全称", FULL_NAME), ("软件简称", SHORT_NAME), ("版本号", VERSION), ("软件分类", "应用软件"), ("开发完成日期", COMPLETION_DATE), ("发表状态", "已发表"), ("首次发表日期", FIRST_PUBLICATION_DATE), ("首次发表方式", "GitHub公开仓库：https://github.com/changfeng-01/edex"), ("首次发表城市", "【首次发表城市】")]),
        ("二、开发方式与权利状况", [("开发方式", "合作开发"), ("权利取得方式", "原始取得"), ("权利范围", "全部权利"), ("著作权人", "实际开发成员共同申请；一名实际开发成员一行"), ("团队名称", "仅作项目协作标识，不作为无主体资格的单一著作权人"), ("申请代表", "【申请代表姓名及授权依据】")]),
        ("三、共同申请人信息", [("共同申请人1", "【姓名】｜【证件类型及号码】｜【地址】｜【电话】｜【邮箱】"), ("共同申请人2", "【姓名】｜【证件类型及号码】｜【地址】｜【电话】｜【邮箱】"), ("共同申请人3", "【姓名】｜【证件类型及号码】｜【地址】｜【电话】｜【邮箱】"), ("可扩展行", "按实际开发成员人数复制；非开发参赛成员和指导教师不自动列入"), ("签章核对", "正式申请表、合作合同与确认书中的姓名和顺序一致")]),
        ("四、软硬件环境与源码规模", [("硬件环境", "通用x86-64个人计算机或服务器；建议8 GB及以上内存、2 GB及以上可用磁盘空间"), ("软件环境", "Windows 10/11或主流Linux；Python 3.11；Node.js；现代浏览器；SQLite"), ("开发语言", "Python、TypeScript/TSX、CSS"), ("主要框架", "FastAPI、SQLAlchemy、React、Vite"), ("源程序量", f"自研源码约{source_stats['lines']}行；{source_stats['files']}个纳入清单文件；鉴别材料3000行")]),
        ("五、主要功能和技术特点", [("主要功能", "项目化管理电路仿真CSV；创建设计版本；约束分析；生成候选；人工审批；导出仿真任务；回填结果；版本比较；查看报告。"), ("技术特点", "以版本、分析运行、候选、实验和仿真任务为领域对象；前后端分离；结构化状态机；工件哈希追溯；显式证据边界。"), ("数据边界", BOUNDARIES[0]), ("工程效力", BOUNDARIES[1]), ("重仿真要求", BOUNDARIES[2]), ("限制声明", "不宣称大量实验验证、流片、物理实测或保证优化效果。")]),
        ("六、附件与正式生成检查", [("程序鉴别材料", "04源程序鉴别材料，60页，每页50行，前后各连续30页"), ("文档鉴别材料", "03操作手册全文，18页，不足60页按全文提交"), ("权属文件", "05合作合同和05A确认书，补齐并逐人签署"), ("正式申请表", "必须在官方系统录入并在线生成；本底稿不是官方表格"), ("禁止事项", "不得伪造、复制或拼接官方二维码、流水号、受理章或签章"), ("提交前", "核对申请人身份、首次发表城市、申请代表、签订地点和所有签字日期")]),
    ]
    for index, (title, rows) in enumerate(page_rows, 1):
        _begin_template_page(doc, index, 6, "软件著作权登记申请表填报底稿", profile="application")
        _add_page_title(doc, title, "非官方申请表，仅用于系统录入；正式表须在线生成")
        add_kv_table(doc, rows, widths=(4.2, 13.0))
        if index in (2, 5, 6):
            add_boundary(doc)
    path = output / "01_软件著作权登记申请表填报底稿.docx"
    save_doc(doc, path)
    return path


def build_template_design(output: Path, source_stats: dict) -> Path:
    doc = _new_template_doc()
    _begin_template_page(doc, 1, 12, "业务理解与软件设计说明书", cover=True)
    _add_template_cover(doc, "业务理解与软件设计说明书", "CircuitPilot V1.0 模板对齐版", "内部支撑材料")
    sections = [
        ("目录与文档控制", ["第1章 软件概述与业务背景", "第2章 用户角色与业务闭环", "第3章 总体架构与模块职责", "第4章 数据模型、数据流与接口", "第5章 状态机、异常与证据链", "第6章 安全边界、部署和验收"], [("基线提交", BASELINE_COMMIT), ("源码范围", f"{source_stats['files']}个文件、{source_stats['lines']}行"), ("材料性质", "内部支撑，不替代正式操作手册或官方申请表")]),
        ("软件概述与业务背景", ["芯智调参服务于电路工程人员在仿真结果、设计版本、约束问题和候选参数之间建立可复核闭环。", "传统分散在CSV、脚本和人工记录中的过程，被整理为项目、设计版本、分析运行、候选、实验、任务和比较等对象。", "软件输出是下一轮工程判断的输入，不把算法建议包装成已验证结论。"], []),
        ("用户角色与核心闭环", ["工程操作员：创建项目、上传仿真CSV、确认字段、启动分析、审批候选、导出任务和回填结果。", "评审人员：查看证据、约束问题、候选理由、版本比较和报告。", "共同开发成员：维护接口、状态规则、测试、部署脚本和申报材料。", "闭环顺序：工作区→项目→设计版本→输入预览→分析→候选审批→仿真任务→外部仿真→结果回填→比较。"], []),
        ("总体架构", ["表现层采用React和TypeScript提供项目工作台；Product API采用FastAPI公开领域操作；服务层实现项目、输入、分析、候选、实验、任务与比较逻辑；仓储层使用SQLAlchemy和SQLite保存结构化状态；工件层保存CSV、JSON、YAML与报告。", "所有跨层操作围绕稳定标识符和显式状态进行，便于审计、重放和问题定位。"], [("前端", "React / TypeScript / Vite"), ("后端", "Python / FastAPI / SQLAlchemy"), ("存储", "SQLite + 本地工件目录"), ("外部系统", "电路仿真器由用户在软件之外运行")]),
        ("模块职责", ["工作区与项目模块负责业务隔离和导航；设计版本模块保存参数、输入工件和来源；分析模块执行CSV读取、指标计算、约束检查和报告形成；候选模块记录建议参数、理由和审批状态；实验与仿真任务模块组织导出及回填；比较模块对两个版本的指标和约束结果进行并列展示。"], [("只读仪表盘", "提供历史工件和展示接口"), ("上传分析服务", "接收输入、预览字段并生成分析入口"), ("评估/优化内核", "提供指标、规则、PIA/LLSO及多智能体协作能力")]),
        ("核心数据模型与数据流", ["工作区包含项目；项目包含设计版本；设计版本关联输入工件和分析运行；分析运行产生问题、证据和候选；候选经审批后形成实验和仿真任务；回填后的新版本可与基线版本比较。", "CSV只作为仿真数据工件进入系统，参数建议不会直接修改物理电路。"], [("输入", "仿真CSV、约束、参数与项目元数据"), ("处理", "校验、解析、指标计算、规则判断、候选组织"), ("输出", "证据、候选、任务包、比较结果和报告")]),
        ("API与接口契约", ["接口按工作区、项目、设计版本、分析运行、候选、实验、仿真任务和比较资源组织。创建或变更操作返回资源标识和当前状态；读取操作返回结构化JSON；错误采用明确状态码与错误对象。", "前端不绕过服务层直接操作数据库或工件目录，外部仿真器也不被伪装为软件内部能力。"], [("示例", "POST /api/v1/workspaces/{id}/projects"), ("示例", "POST /api/v1/design-versions/{id}/analysis-runs"), ("示例", "POST /api/v1/candidates/{id}/approve"), ("示例", "GET /api/v1/comparisons/{id}")]),
        ("领域状态机", ["设计版本在创建、就绪和可分析状态之间转换；分析运行经历排队、运行、成功或失败；候选经历提出、批准/拒绝、导出、评估；仿真任务经历准备、导出、等待外部结果、回填和完成。", "状态转换由服务层校验，非法重复操作或顺序错误返回冲突，不通过界面文字掩盖。"], [("关键规则", "候选批准不等于工程有效"), ("关键规则", "结果回填后仍需重新分析和比较"), ("关键规则", "neutral或未改善结果保持原样展示")]),
        ("异常处理", ["输入缺失、CSV字段不匹配、数值解析失败、资源不存在、状态冲突、工件哈希不一致和外部结果缺失均形成可定位错误。", "界面向用户展示可执行修复提示；服务日志保留技术上下文，但正式截图和申报材料不带本地绝对路径、密钥或调试堆栈。"], [("4xx", "输入、资源或状态问题"), ("5xx", "未预期服务异常"), ("恢复", "修正工件或状态后重新执行，不伪造成功结果")]),
        ("证据链与安全边界", ["每个源文件、提交材料和输入工件均通过路径、行号、资源标识或SHA-256形成追溯。测试夹具只用于展示流程，不作为工程效果证据。", "系统不在材料中保存环境变量、密钥、申请人身份证件或未脱敏私有数据；正式签署信息由团队在线下或官方系统补齐。"], [("数据来源", BOUNDARIES[0]), ("工程效力", BOUNDARIES[1]), ("后续要求", BOUNDARIES[2])]),
        ("部署、验收、版本与限制", ["部署由Python后端、React前端、SQLite数据库和本地工件目录组成；开发环境可分别启动API与Vite，生产构建由静态资源和后端服务共同交付。", "验收覆盖后端pytest、前端测试与生产构建、确定性产品演示、文档页数/哈希/占位符核对以及DOCX、PDF逐页渲染检查。", f"本材料对应{VERSION}，申报基线为{BASELINE_COMMIT}。本轮仅调整版本元数据、材料生成和验证工具，不改变Product API、CLI、数据库模型或前端业务行为。", "当前材料证明软件结构、代码与可运行流程，不证明大量实验验证、流片、芯片实测或优化效果保证。"], [("运行环境", "本机或受控内网；现代Chromium浏览器"), ("备份", "SQLite数据库与工件目录一并备份"), ("完成/发表日期", f"{COMPLETION_DATE} / {FIRST_PUBLICATION_DATE}"), ("申请口径", "实际开发成员共同申请；合作开发；原始取得；全部权利")]),
    ]
    for physical, (title, paragraphs, rows) in enumerate(sections, 2):
        _begin_template_page(doc, physical, 12, "业务理解与软件设计说明书", body_page=physical - 1, body_total=11)
        _add_page_title(doc, title)
        _add_body_paragraphs(doc, paragraphs)
        if rows:
            add_kv_table(doc, rows, widths=(4.0, 12.0))
        if physical in (3, 10, 12):
            add_boundary(doc)
    path = output / "02_业务理解与软件设计说明书.docx"
    save_doc(doc, path)
    return path


def build_template_manual(output: Path, screenshot_dir: Path) -> Path:
    doc = _new_template_doc()
    _begin_template_page(doc, 1, 18, "软件操作手册暨文档鉴别材料", cover=True)
    _add_template_cover(doc, "软件操作手册暨文档鉴别材料", "完整文档；正文17页连续编号", "正式文档鉴别材料候选")
    manual_pages = [
        ("目录", ["1 软件概述与边界　2 环境与安装　3 启动与工作区　4 创建项目　5 上传仿真数据", "6 设计版本预览　7 分析　8 候选审批　9 仿真任务导出　10 外部仿真与回填", "11 结果版本　12 结果分析　13 版本比较　14 报告查看　15 异常处理　16 安全说明"], None),
        ("软件概述与证据边界", ["软件用于管理电路仿真CSV、分析约束、组织候选参数和重仿真闭环。候选不是已验证结论，必须由用户在外部仿真器复核。", *BOUNDARIES, "本手册截图均为test_only确定性演示数据，不构成实验、物理实测或流片证明。"], "01_public_demo.png"),
        ("环境、安装与启动", ["准备Python 3.11、Node.js、现代浏览器和可写工件目录；安装后端及前端依赖。", "启动Product API，再启动前端开发或生产服务；浏览器访问工作区页面。", "生产部署时应设置受控数据库和工件目录，并通过组织现有网络边界提供访问。"], None),
        ("创建项目", ["在工作区项目页选择“创建项目”，填写项目名称、说明和目标约束；确认后系统生成项目标识并进入项目工作台。", "项目名称仅用于业务区分，不能代替申请人或著作权人信息。"], "03_create_project.png"),
        ("上传仿真数据", ["进入上传工作区，选择CSV和必要的映射/约束文件；先预览字段和记录数，再提交分析。", "上传文件应来自真实仿真过程；test_only夹具只用于演示页面流程。"], "12_upload_workspace.png"),
        ("设计版本预览", ["打开基线设计版本，核对版本名称、参数、输入工件、来源和状态。", "只有输入工件可读且状态允许时，才继续创建分析运行。"], "05_baseline_design_version.png"),
        ("运行分析", ["在设计版本页面发起分析，系统解析CSV、计算指标、检查约束并形成问题、证据与候选。", "页面显示的结论受仿真输入和约束配置限制，不扩展为物理工程保证。"], "06_baseline_analysis.png"),
        ("候选审批", ["查看候选参数、修改理由、关联问题和证据；工程人员可批准或拒绝。", "批准只表示同意进入下一轮仿真，不表示候选已改善或已通过实验。"], "07_candidate_approval.png"),
        ("仿真任务导出", ["批准候选后创建仿真任务，下载任务包并记录任务标识、版本和预期回填格式。", "任务状态用于组织外部工作，不自动调用或替代第三方电路仿真器。"], "08_simulation_job.png"),
        ("外部仿真与结果回填", ["在外部仿真器中按任务包应用候选参数，运行仿真并导出结果CSV。", "回到系统按任务要求上传结果，创建新的设计版本；若文件不符合契约，修正后重新提交。", BOUNDARIES[2]], "08_simulation_job.png"),
        ("查看结果设计版本", ["结果回填后打开新设计版本，核对来源任务、参数和新输入工件。", "新版本与基线版本并存，避免覆盖原始证据。"], "09_result_design_version.png"),
        ("重新分析结果版本", ["对结果版本再次运行分析，查看指标、约束问题和证据。", "若结果为neutral、未改善或失败，应原样保留，不改写为成功。"], "10_result_analysis.png"),
        ("版本比较", ["选择基线版本和结果版本进行比较，查看指标变化、约束变化和总体判定。", "比较仅描述相同口径下的仿真数据差异。"], "11_version_comparison.png"),
        ("报告与公开演示页", ["报告页面汇总项目、版本、问题、候选、任务和边界声明；导出前核对是否含本地路径或调试信息。", "公开演示页可用于说明软件工作流，不可作为真实性能宣传。"], "01_public_demo.png"),
        ("异常处理", ["WORKSPACE_NOT_FOUND：确认工作区标识或返回工作区列表重新选择。", "CSV_SCHEMA_INVALID：核对列名、编码和数值格式。", "STATE_CONFLICT：按页面提示完成前置状态，避免重复批准或回填。", "ARTIFACT_MISSING：恢复工件或重新上传；不得伪造成功状态。"], None),
        ("安全与隐私说明", ["截图、导出包和申报材料不得包含密钥、环境变量、身份证号、未脱敏私有数据和本地绝对路径。", "共同申请人的身份和签章仅在官方系统或经授权的线下文件中补齐。", "第三方仿真器、PDK和依赖按其许可证使用。"], None),
        ("停止、备份与验收", ["停止服务前确认分析和回填任务已结束；备份SQLite数据库和工件目录。", "恢复时保持数据库与工件版本一致，并抽查项目、版本、分析和比较页面。", "本手册共18个PDF页面：封面不编号，正文1—17页连续编号；因不足60页，按完整文档提交。"], None),
    ]
    for physical, (title, paragraphs, screenshot) in enumerate(manual_pages, 2):
        _begin_template_page(doc, physical, 18, "软件操作手册暨文档鉴别材料", body_page=physical - 1, body_total=17)
        _add_page_title(doc, title)
        _add_body_paragraphs(doc, paragraphs)
        if screenshot:
            image_path = screenshot_dir / screenshot
            if not image_path.is_file():
                raise FileNotFoundError(f"手册截图不存在：{image_path}")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(5)
            p.add_run().add_picture(str(image_path), width=Cm(15.8))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(f"图{physical - 1}　test_only 真实运行界面（1440×900）")
            set_run_font(r, east_asia="宋体", size=8, color=GRAY)
        if physical in (3, 11, 18):
            add_boundary(doc)
    path = output / "03_软件操作手册暨文档鉴别材料.docx"
    save_doc(doc, path)
    return path


def build_template_contract(output: Path) -> Path:
    doc = _new_template_doc()
    _begin_template_page(doc, 1, 16, "技术开发（合作）合同", cover=True)
    _add_template_cover(doc, "技术开发（合作）合同", "依据合作合同模板形成的项目底稿", "权属证明；签署后使用")
    clause_pairs = [
        ("第一条　合同主体与项目", "甲方/合作方一：【姓名】、【证件号码】、【地址】、【电话】；乙方/合作方二：【姓名】、【证件号码】、【地址】、【电话】；其他合作方按实际开发成员增列。项目名称为“芯智调参：基于仿真数据的电路参数智能推荐系统V1.0”。", "第二条　合作目标", "共同完成电路仿真数据管理、约束分析、候选参数管理、仿真任务导出、结果回填、版本比较与报告查看的软件，并形成源码、界面、测试和文档。"),
        ("第三条　技术内容", "合作内容包括Python后端、React前端、数据模型、工件追溯、评估与优化内核、测试及申报文档；不包括第三方仿真器、PDK和第三方组件本身。", "第四条　技术方法和路线", "采用前后端分离、领域服务、结构化状态机、SQLite与工件目录存储；以仿真CSV为输入，通过分析、候选审批、外部重仿真和版本比较形成闭环。"),
        ("第五条　计划进度", "各方根据版本计划完成需求、实现、测试和材料；V1.0开发完成日期确认为2026年07月15日。后续版本另行书面确认，不当然改变本合同对V1.0的约定。", "第六条　组织方式", "各方以版本库、问题清单、评审记录和测试结果协同；重大架构、权利处分和对外发布事项由共同开发成员书面决定。"),
        ("第七条　成员分工", "各方实际贡献在本合同附件《共同开发与著作权归属确认书》中逐人填写，并以提交记录、评审记录、文档或测试证据索引佐证。", "第八条　开发成本", "除各方另有书面约定外，各方自行承担本人投入的设备、时间、差旅和日常开发成本；共同采购或外部支出须事先书面同意。"),
        ("第九条　技术资料提供", "各方应提供本人合法持有且完成任务所需的资料，不得将无权使用的源码、数据、密钥或私有材料纳入项目。", "第十条　协作与通知", "影响交付、权属或安全的事项应及时在可追溯渠道通知其他合作方；联系方式变化应书面更新。"),
        ("第十一条　成果交付", "交付物包括版本库源码、构建配置、测试、演示工件、操作手册、设计说明、源程序鉴别材料和相关清单。交付以指定提交和SHA-256记录为准。", "第十二条　验收标准", "后端测试、前端测试和生产构建通过；确定性产品流程可运行；材料名称、版本、日期、页数、边界与哈希检查通过。"),
        ("第十三条　验收方式", "各方共同查验版本库、测试输出、演示流程和申报材料；发现问题形成清单并由相应贡献人修正。未提出异议不当然免除权属瑕疵责任。", "第十四条　风险承担", "因外部仿真器、PDK、第三方服务、网络或硬件产生的风险由实际使用方按许可和环境承担；软件不保证候选参数必然改善。"),
        ("第十五条　保密范围", "未公开源码、私有数据、密钥、申请人证件信息、联系方式和合作方书面标注的资料属于保密信息；依法公开或已合法公开的信息除外。", "第十六条　保密义务", "接收方仅为项目和登记目的使用保密信息，采取合理访问控制；未经信息主体和其他相关方书面同意不得对外披露。"),
        ("第十七条　知识产权归属", "各实际开发成员对V1.0团队原创代码、原创界面编排和原创文档共同享有软件著作权；登记采用合作开发、原始取得、全部权利。", "第十八条　第三方权利", "Python、FastAPI、SQLAlchemy、React等第三方组件按各自许可证使用，不纳入共同著作权主张；引用资料不得改变原权利归属。"),
        ("第十九条　申请与代表", "各方同意共同申请软件著作权，并授权【申请代表】在授权范围内办理材料录入、提交、补正和证书领取；代表不得单方处分共有权利。", "第二十条　署名与对外表述", "团队名称可作为项目标识使用，但不作为无主体资格的单一著作权人；非开发参赛成员和指导教师不自动列入共同著作权人。"),
        ("第二十一条　收益分配", "许可、转让或其他处置产生的收益和费用由各方另行书面约定；未约定前不得由任何一方单独占有或处分。", "第二十二条　重大处分", "著作权转让、独占许可、质押、放弃权利、重大开源或其他重大处分必须经全体共同著作权人书面同意。"),
        ("第二十三条　后续改进", "个人独立完成且可与V1.0区分的后续改进，其权属另行认定；基于共有代码形成且不可分割的改进，应由相关方书面约定。", "第二十四条　发表与论文", "代码公开、论文发表和竞赛展示应遵守保密与共同署名约定。软件登记不当然限制论文发表，但不得披露未获授权的身份、私有数据或第三方机密。"),
        ("第二十五条　违约责任", "一方擅自处分共有权利、引入侵权材料、泄露保密信息或提供虚假申请信息的，应停止行为、协助补救并承担由此造成的实际损失。", "第二十六条　不可抗力", "因不可预见、不可避免且不可克服事件导致迟延的，受影响方应及时通知并提供合理证明，各方协商调整计划。"),
        ("第二十七条　争议解决", "争议先由各方协商；协商不成，向【约定有管辖权的人民法院或仲裁机构】处理。适用中华人民共和国法律。", "第二十八条　合同生效与变更", "本合同经全部合作方签字或盖章后生效；补充、变更、附件和授权须采用可核验书面形式，并与本合同具有同等效力。"),
    ]
    for physical, clauses in enumerate(clause_pairs, 2):
        _begin_template_page(doc, physical, 16, "技术开发（合作）合同")
        _add_page_title(doc, f"合同条款（{physical - 1}）")
        for title, text in ((clauses[0], clauses[1]), (clauses[2], clauses[3])):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            r = p.add_run(title)
            set_run_font(r, east_asia="黑体", size=12, bold=True)
            _add_body_paragraphs(doc, [text])
        if physical in (3, 10):
            add_boundary(doc)
    _begin_template_page(doc, 16, 16, "技术开发（合作）合同")
    _add_page_title(doc, "第二十九条　签署页与附件")
    _add_body_paragraphs(doc, ["本页为签署页，无正文。合同编号、签订地点、身份信息、联系方式和签署日期须由各方核实后填写。附件《共同开发与著作权归属确认书》是本合同组成部分。"])
    add_kv_table(doc, [
        ("合同编号", "【合同编号】"), ("签订地点", "【签订地点】"), ("合作方一", "【姓名】　签字/盖章：【签章】　日期：【年/月/日】"),
        ("合作方二", "【姓名】　签字/盖章：【签章】　日期：【年/月/日】"), ("合作方三", "【姓名】　签字/盖章：【签章】　日期：【年/月/日】"),
        ("其他合作方", "按实际开发成员人数增列并逐人签署"), ("附件", "05A_共同开发与著作权归属确认书"),
    ], widths=(4.0, 12.0))
    path = output / "05_技术开发（合作）合同.docx"
    save_doc(doc, path)
    return path


def build_template_confirmation(output: Path) -> Path:
    doc = _new_template_doc()
    _begin_template_page(doc, 1, 4, "共同开发与著作权归属确认书", cover=True)
    _add_template_cover(doc, "共同开发与著作权归属确认书", "技术开发（合作）合同附件", "权属证明；逐人签署")
    page_content = [
        ("一、共同开发成员与实际贡献", [
            ("成员1", "【姓名】｜【证件号码】｜【实际开发贡献】｜【提交/评审/文档证据索引】"),
            ("成员2", "【姓名】｜【证件号码】｜【实际开发贡献】｜【提交/评审/文档证据索引】"),
            ("成员3", "【姓名】｜【证件号码】｜【实际开发贡献】｜【提交/评审/文档证据索引】"),
            ("增列规则", "一名实际开发成员一行；非开发参赛成员和指导教师不自动列入"),
        ]),
        ("二、共同权属与申请授权", [
            ("权属", "各实际开发成员对V1.0团队原创代码、原创界面编排和原创文档共同享有著作权"),
            ("申请口径", "合作开发、原始取得、全部权利、共同申请"),
            ("申请代表", "授权【申请代表姓名】办理系统录入、提交、补正和证书领取；不得单方处分共有权利"),
            ("处分", "转让、独占许可、质押、放弃和重大开源须全体共同著作权人书面同意"),
            ("第三方", "不主张第三方组件、仿真器、PDK或他人资料的著作权"),
        ]),
        ("三、确认与签署", [
            ("成员1", "【姓名】　签字：【签字】　日期：【年/月/日】"),
            ("成员2", "【姓名】　签字：【签字】　日期：【年/月/日】"),
            ("成员3", "【姓名】　签字：【签字】　日期：【年/月/日】"),
            ("其他成员", "按实际开发成员人数增列并逐人签署"),
            ("共同确认", "所填贡献真实、权属无争议、申请信息一致；如有变更应由全体成员书面确认"),
        ]),
    ]
    for physical, (title, rows) in enumerate(page_content, 2):
        _begin_template_page(doc, physical, 4, "共同开发与著作权归属确认书")
        _add_page_title(doc, title)
        add_kv_table(doc, rows, widths=(4.0, 12.0))
        if physical in (2, 3):
            add_boundary(doc)
    path = output / "05A_共同开发与著作权归属确认书.docx"
    save_doc(doc, path)
    return path


def build_template_third_party(output: Path) -> Path:
    doc = _new_template_doc()
    _begin_template_page(doc, 1, 3, "第三方组件与权利边界说明", cover=True)
    _add_template_cover(doc, "第三方组件与权利边界说明", "依赖、许可证与原创范围核对", "内部支撑材料")
    _begin_template_page(doc, 2, 3, "第三方组件与权利边界说明", body_page=1, body_total=2)
    _add_page_title(doc, "主要第三方组件")
    add_table(doc, ["组件", "许可证", "用途", "权利边界"], [
        ("Python", "PSF", "后端运行时", "不主张运行时著作权"), ("FastAPI", "MIT", "HTTP API框架", "仅按许可调用"),
        ("SQLAlchemy", "MIT", "数据库访问", "仅按许可调用"), ("NumPy / pandas", "BSD-3-Clause", "数值与表格处理", "仅按许可调用"),
        ("React / React DOM", "MIT", "前端界面框架", "仅按许可调用"), ("React Router", "MIT", "前端路由", "仅按许可调用"),
        ("Recharts", "MIT", "图表展示", "仅按许可调用"), ("Vite / TypeScript", "MIT / Apache-2.0", "构建与语言工具", "不纳入登记主张"),
        ("Vitest", "MIT", "前端测试", "测试工具不纳入运行成果主张"),
    ], font_size=7.8)
    _begin_template_page(doc, 3, 3, "第三方组件与权利边界说明", body_page=2, body_total=2)
    _add_page_title(doc, "原创范围、排除项与合规核对")
    add_kv_table(doc, [
        ("登记范围", "团队原创Python/TypeScript/TSX/CSS源码、原创领域结构、界面编排和原创文档"),
        ("不主张", "第三方组件、框架、运行时、仿真器、PDK、开源示例、字体和他人资料"),
        ("源码排除", "tests、node_modules、锁文件、构建输出、环境变量、私有数据、自动生成样板和本材料生成工具"),
        ("许可证核对", "正式发布前以锁文件和上游许可证原文复核版本、版权声明和再分发条件"),
        ("界面截图", "真实运行界面由原创业务代码编排；框架提供通用能力，不改变第三方权属"),
        ("外部仿真", "系统组织任务和回填，不主张外部仿真器或其输出格式的著作权"),
    ], widths=(4.0, 12.0))
    add_boundary(doc)
    path = output / "06_第三方组件与权利边界说明.docx"
    save_doc(doc, path)
    return path


def copy_support_screenshots(output: Path) -> Path:
    target = output / "internal_support" / "screenshots_test_only"
    target.mkdir(parents=True, exist_ok=True)
    for path in sorted(SCREENSHOT_SOURCE.glob("*.png")):
        if not re.match(r"^\d{2}_.+\.png$", path.name):
            continue
        shutil.copy2(path, target / path.name)
    readme = target / "README.txt"
    readme.write_text(
        "本目录截图由 Playwright 在 1440×900 视口下采集自真实运行界面。\n"
        "数据为确定性 test_only 演示夹具，仅用于软件工作流和界面说明。\n"
        "不得解释为真实实验、物理实测、芯片验证或流片证明。\n",
        encoding="utf-8",
    )
    return target


def write_template_alignment_audit(output: Path, entries: list[dict]) -> Path:
    support = output / "internal_support"
    support.mkdir(parents=True, exist_ok=True)
    payload = {
        "template_authority": TEMPLATE_ALIGNMENT,
        "facts_source": "CircuitPilot repository and deterministic test_only product demo",
        "approved_example_use": "visual acceptance only; no names, seals, QR codes or software facts copied",
        "document_specs": DOCUMENT_SPECS,
        "source_inventory": {
            "file_count": len(entries),
            "line_count": sum(item["line_count"] for item in entries),
            "group_counts": dict(Counter(item["source_group"] for item in entries)),
        },
        "evidence_boundary": list(BOUNDARIES),
        "formal_submission_recommendation": {
            "official_form": "generate online from the authority system",
            "document_identification": "submit the complete 18-page operation manual",
            "program_identification": "submit the 60-page source-code material",
            "rights_documents": "submit signed contract/confirmation when required",
        },
    }
    path = support / "template_alignment_audit.json"
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    traceability = support / "design_traceability.csv"
    grouped: dict[str, list[dict]] = {}
    for entry in entries:
        grouped.setdefault(entry["source_group"], []).append(entry)
    with traceability.open("w", newline="", encoding="utf-8-sig") as file:
        writer = csv.DictWriter(
            file,
            fieldnames=["设计职责", "源码分组", "代表文件", "文件数", "验证依据"],
        )
        writer.writeheader()
        for group, group_entries in grouped.items():
            writer.writerow(
                {
                    "设计职责": {
                        "产品 API": "资源接口、错误契约与状态入口",
                        "产品服务": "项目、版本、分析、候选、任务和比较领域逻辑",
                        "只读仪表盘 API": "历史工件与只读展示接口",
                        "上传分析服务": "输入预览、上传与分析入口",
                        "前端界面": "项目工作台与闭环操作界面",
                        "工程脚本": "演示、迁移、校验和运行编排",
                        "电路分析": "电路指标与约束分析",
                        "PIA 与候选生成": "候选参数组织与推荐依据",
                        "LLSO 优化": "优化搜索与候选生成",
                        "多智能体协作": "分析协作与结构化汇总",
                        "评估内核与命令行": "数据读取、评估、诊断、报告与CLI",
                    }.get(group, "其他原创实现"),
                    "源码分组": group,
                    "代表文件": "；".join(item["path"] for item in group_entries[:5]),
                    "文件数": len(group_entries),
                    "验证依据": "source_manifest.csv、source_line_manifest.csv、pytest/前端测试与产品演示",
                }
            )
    return path


def write_snapshot_manifest(
    output: Path,
    source_entries: list[dict],
    source_info: dict,
    generated_docs: list[Path],
) -> Path:
    screenshot_files = [
        path for path in sorted(SCREENSHOT_SOURCE.glob("*.png"))
        if re.match(r"^\d{2}_.+\.png$", path.name)
    ]
    manifest = {
        "schema_version": MANIFEST_SCHEMA_VERSION,
        "software": {
            "full_name": FULL_NAME,
            "short_name": SHORT_NAME,
            "version": VERSION,
            "package_version": PACKAGE_VERSION,
            "completion_date": "2026-07-15",
            "publication_status": "已发表",
            "first_publication_date": "2026-07-15",
            "first_publication_method": "GitHub公开仓库",
            "first_publication_url": "https://github.com/changfeng-01/edex",
            "first_publication_city": "【首次发表城市】",
        },
        "application": {
            "applicants": "实际开发成员共同申请",
            "development_mode": "合作开发",
            "rights_acquisition": "原始取得",
            "rights_scope": "全部权利",
        },
        "git": {
            "baseline_commit": BASELINE_COMMIT,
            "branch": BRANCH,
            "current_head": run_text(["git", "rev-parse", "HEAD"]),
        },
        "template_alignment": TEMPLATE_ALIGNMENT,
        "materials": {
            stem: {
                "expected_pdf_pages": spec["pdf_pages"],
                "submission_role": spec["submission_role"],
                "docx": f"{stem}.docx",
                "pdf": f"{stem}.pdf",
                "formal_submission_recommended": stem in {
                    "03_软件操作手册暨文档鉴别材料",
                    "04_源程序鉴别材料_前30页后30页",
                    "05_技术开发（合作）合同",
                    "05A_共同开发与著作权归属确认书",
                },
            }
            for stem, spec in DOCUMENT_SPECS.items()
        },
        "evidence_boundary": {
            "data_source": "real_simulation_csv",
            "engineering_validity": "simulation_only",
            "must_resimulate": True,
        },
        "source": {
            "file_count": len(source_entries),
            "line_count": sum(item["line_count"] for item in source_entries),
            "manifest": "source_manifest.csv",
            "extensions": [".py", ".ts", ".tsx", ".css"],
            "excluded": [
                "tests and test files",
                "third-party dependencies and node_modules",
                "lock files",
                "outputs and generated artifacts",
                "environment files and private data",
                "automatic boilerplate and material-generation tools",
            ],
            "identification_material": source_info,
        },
        "document_identification_material": {
            "path": "03_软件操作手册暨文档鉴别材料.docx",
            "expected_pdf_pages": 18,
            "complete_document": True,
            "under_60_pages": True,
            "numbering": "cover unnumbered; body pages 1-17 continuous",
            "formal_submission_recommended": True,
        },
        "screenshots": {
            "source": "Playwright CLI",
            "viewport": "1440x900",
            "fixture": "test_only deterministic product demo",
            "files": [
                {"path": p.name, "sha256": sha256(p), "bytes": p.stat().st_size}
                for p in screenshot_files
            ],
        },
        "generated_docx": [
            {"path": p.name, "sha256": sha256(p), "bytes": p.stat().st_size}
            for p in generated_docs
        ],
        "validation": {
            "status": "pending_post_render_validation",
            "required": [
                "pytest",
                "frontend tests and production build",
                "product demo",
                "DOCX to PDF export",
                "page count and fixed-line checks",
                "PDF re-render visual inspection",
                "metadata and placeholder audit",
            ],
            "allowed_placeholders": [
                "applicant identity and contact details",
                "individual contribution and evidence index",
                "first publication city",
                "contract number and signing place",
                "application representative",
                "signature/seal and signing date",
                "dispute resolution venue",
            ],
        },
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }
    path = output / "snapshot_manifest.json"
    path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def main() -> int:
    output = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else DEFAULT_OUTPUT
    ensure_safe_output_path(output)
    output.mkdir(parents=True, exist_ok=True)
    if not SCREENSHOT_SOURCE.is_dir():
        raise FileNotFoundError(f"截图目录不存在：{SCREENSHOT_SOURCE}")
    entries = list_source_files()
    if not entries:
        raise RuntimeError("没有找到符合口径的自研源文件。")
    write_source_manifest(output, entries)
    source_stats = {
        "files": len(entries),
        "lines": sum(item["line_count"] for item in entries),
        "entries": entries,
    }
    support_screenshots = copy_support_screenshots(output)
    write_template_alignment_audit(output, entries)
    docs: list[Path] = []
    docs.append(build_template_catalog(output, source_stats))
    docs.append(build_template_application(output, source_stats))
    docs.append(build_template_design(output, source_stats))
    docs.append(build_template_manual(output, support_screenshots))
    source_info = build_source_doc(output, entries)
    docs.append(output / source_info["path"])
    docs.append(build_template_contract(output))
    docs.append(build_template_confirmation(output))
    docs.append(build_template_third_party(output))
    manifest = write_snapshot_manifest(output, entries, source_info, docs)
    print(
        json.dumps(
            {
                "output": str(output),
                "source_files": source_stats["files"],
                "source_lines": source_stats["lines"],
                "docx": [p.name for p in docs],
                "manifest": manifest.name,
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
